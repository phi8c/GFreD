import os
from docx import Document
import openpyxl
from openpyxl import Workbook
import zipfile
from openpyxl.styles import Font, Alignment
from openpyxl.utils import get_column_letter
import tempfile
import shutil
from openpyxl import load_workbook
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from flask import send_file, request, current_app
import io, requests

from controllers.db import get_db_connection

from app.global_config import UPLOAD_FOLDER

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib import colors
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics

# Đăng ký font Unicode có sẵn
#pdfmetrics.registerFont(UnicodeCIDFont('HeiseiMin-W3'))
pdfmetrics.registerFont(TTFont("Roboto", "fonts/Roboto-Regular.ttf"))
from io import BytesIO
from models.exam_model import get_exam_codes_by_set,  get_exam_questions_by_code, get_original_questions_by_exam_set, get_original_questions_thpt2025, get_original_questions_thpt2025, is_thpt2025_exam_set, get_exam_mcq_questions_by_code, get_exam_code_id_by_code, get_questions_by_exam_set, insert_mcq_question


styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name='RobotoNormal', fontName='Roboto', fontSize=11, leading=14))
styles.add(ParagraphStyle(name='RobotoTitle', fontName='Roboto', fontSize=16, leading=20, alignment=1))  # center
styles.add(ParagraphStyle(name='RobotoHeading2', fontName='Roboto', fontSize=13, leading=18, spaceAfter=6))



def create_export_folder(set_id):
    folder_path = f"exported_exams/set_{set_id}"
    os.makedirs(folder_path, exist_ok=True)
    return folder_path
def export_exam_to_docx(exam_code, questions, folder_path):
    doc = Document()
    doc.add_heading(f"ĐỀ THI - MÃ ĐỀ {exam_code}", level=1)

    for idx, q in enumerate(questions, 1):
        doc.add_paragraph(f"{idx}. {q['question_text']}")
        doc.add_paragraph(f"A. {q['answer_a']}")
        doc.add_paragraph(f"B. {q['answer_b']}")
        doc.add_paragraph(f"C. {q['answer_c']}")
        doc.add_paragraph(f"D. {q['answer_d']}")
        doc.add_paragraph("\n")

    doc.save(os.path.join(folder_path, f"de_{exam_code}.docx"))
# new
def export_original_questions_to_docx(questions, folder_path):
    doc = Document()
    doc.add_heading(f"ĐỀ GỐC CHƯA HOÁN VỊ", level=1)

    for idx, q in enumerate(questions, 1):
        doc.add_paragraph(f"{idx}. {q['question_text']}")
        doc.add_paragraph("\n")

    doc.save(os.path.join(folder_path, "de_goc.docx"))




# end

def export_answer_sheet_excel(exam_sets, folder_path):
    wb = Workbook()
    ws = wb.active
    ws.title = "AnswerSheet"

    ws.cell(row=1, column=1, value="STT")

    for col, (exam_code, answers) in enumerate(exam_sets.items(), start=2):
        ws.cell(row=1, column=col, value=f"Mã đề {exam_code}")
        for row, answer in enumerate(answers, start=2):
            if col == 2:
                ws.cell(row=row, column=1, value=row-1)
            ws.cell(row=row, column=col, value=answer)

    wb.save(os.path.join(folder_path, "answer_sheet.xlsx"))
    
# hàm này mới thêm vào


# đây là các hàm mới thêm vào

def export_original_thpt2025_to_docx(questions_by_type, folder_path):
    """
    Xuất đề gốc (chưa hoán vị) theo định dạng THPT 2025: MCQ, TF, SA phân rõ phần.
    `questions_by_type` là dict gồm: {"MCQ": [...], "TF": [...], "SA": [...]}
    """
    doc = Document()
    doc.add_heading("ĐỀ GỐC CHƯA HOÁN VỊ", level=1)

    # Header giống mẫu
    doc.add_paragraph("TRƯỜNG .........................................................", style='Normal')
    doc.add_paragraph("KHOA ..................................................................", style='Normal')
    doc.add_paragraph("MÔN THI: .........................................................", style='Normal')
    doc.add_paragraph("Lớp: ...............................................................", style='Normal')
    doc.add_paragraph("Hình thức thi: ..................     Thời gian: .......... phút", style='Normal')
    doc.add_paragraph("MÃ ĐỀ: .................", style='Normal')
    doc.add_paragraph("Họ và tên: .......................................... MSSV: ...................", style='Normal')
    doc.add_paragraph("")
    doc.add_paragraph("NỘI DUNG ĐỀ THI", style='Normal')

    idx = 1

    # Phần 1: Trắc nghiệm nhiều lựa chọn (MCQ)
    if questions_by_type.get("MCQ"):
        doc.add_paragraph("I. PHẦN TRẮC NGHIỆM NHIỀU LỰA CHỌN", style='Normal')
        for q in questions_by_type["MCQ"]:
            doc.add_paragraph(f"{idx}. {q['question_text']}")
            doc.add_paragraph(f"A. {q['answer_a']}")
            doc.add_paragraph(f"B. {q['answer_b']}")
            doc.add_paragraph(f"C. {q['answer_c']}")
            doc.add_paragraph(f"D. {q['answer_d']}")
            doc.add_paragraph("")
            idx += 1

    # Phần 2: Đúng/Sai (TF)
    if questions_by_type.get("TF"):
        doc.add_paragraph("II. PHẦN ĐÚNG / SAI", style='Normal')
        for q in questions_by_type["TF"]:
            doc.add_paragraph(f"{idx}. {q['question_text']}")
            doc.add_paragraph("☐ Đúng     ☐ Sai")
            doc.add_paragraph("")
            idx += 1

    # Phần 3: Tự luận ngắn (SA)
    if questions_by_type.get("SA"):
        doc.add_paragraph("III. PHẦN TỰ LUẬN", style='Normal')
        for q in questions_by_type["SA"]:
            doc.add_paragraph(f"{idx}. {q['question_text']}")
            doc.add_paragraph("............................................................")
            doc.add_paragraph("............................................................")
            doc.add_paragraph("............................................................")
            doc.add_paragraph("")
            idx += 1

    doc.save(os.path.join(folder_path, "de_goc_thpt2025.docx"))


def export_answer_sheet_thpt2025_excel(answer_data, folder_path):
    """
    Xuất bảng đáp án cho đề thi THPT 2025.
    `answer_data` là dict: { exam_code1: ["A", "C", "Đúng", "Sai", ...], exam_code2: [...] }
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "AnswerSheet"

    ws.cell(row=1, column=1, value="STT")

    for col, (exam_code, answers) in enumerate(answer_data.items(), start=2):
        ws.cell(row=1, column=col, value=f"Mã đề {exam_code}")
        for row, answer in enumerate(answers, start=2):
            if col == 2:
                ws.cell(row=row, column=1, value=row - 1)
            ws.cell(row=row, column=col, value=answer)

    wb.save(os.path.join(folder_path, "answer_sheet_thpt2025.xlsx"))


# kết thúc các hàm mới thêm vào

def export_exam_thpt2025_to_docx(exam_code, questions_by_type, folder_path):
    doc = Document()

    # Thêm phần mở đầu đề thi
    doc.add_paragraph("TRƯỜNG ĐẠI HỌC ..........................................", style='Normal')
    doc.add_paragraph("KHOA: ..................................................", style='Normal')
    doc.add_paragraph("TP. .....................................................", style='Normal')
    doc.add_paragraph("\nĐỀ THI KẾT THÚC HỌC PHẦN", style='Normal')
    doc.add_paragraph("HỌC KỲ ......  NĂM HỌC ...................", style='Normal')
    doc.add_paragraph("\nMÔN THI: ...........................................................", style='Normal')
    doc.add_paragraph("Lớp: ..............................................................", style='Normal')
    doc.add_paragraph("Hình thức thi: ...................     Thời gian: ...... phút. (không kể thời gian phát đề)", style='Normal')
    doc.add_paragraph(f"MÃ ĐỀ: {exam_code}     ☐ Không sử dụng tài liệu     ☐ Được sử dụng tài liệu", style='Normal')
    doc.add_paragraph("              ☐ Nộp lại đề thi             ☐ Không nộp lại đề thi", style='Normal')
    doc.add_paragraph("\nHọ và tên: ..............................................      MSSV: ............................", style='Normal')
    doc.add_paragraph("\nNỘI DUNG ĐỀ THI", style='Normal').bold = True

    # Câu hỏi trắc nghiệm (MCQ)
    mcq_list = questions_by_type.get("mcq", [])
    if mcq_list:
        doc.add_heading("I. TRẮC NGHIỆM (MCQ)", level=2)
        for idx, q in enumerate(mcq_list, 1):
            doc.add_paragraph(f"{idx}. {q['question_text']}")
            doc.add_paragraph(f"A. {q['answer_a']}")
            doc.add_paragraph(f"B. {q['answer_b']}")
            doc.add_paragraph(f"C. {q['answer_c']}")
            doc.add_paragraph(f"D. {q['answer_d']}")

    # Câu hỏi đúng/sai (TF)
    tf_list = questions_by_type.get("tf", [])
    if tf_list:
        doc.add_heading("II. ĐÚNG/SAI (TRUE/FALSE)", level=2)
        for idx, q in enumerate(tf_list, 1):
            doc.add_paragraph(f"{idx}. {q['question_text']} (Đúng/Sai)")

    # Câu hỏi tự luận ngắn (SA)
    sa_list = questions_by_type.get("sa", [])
    if sa_list:
        doc.add_heading("III. TỰ LUẬN NGẮN (SHORT ANSWER)", level=2)
        for idx, q in enumerate(sa_list, 1):
            doc.add_paragraph(f"{idx}. {q['question_text']}")
            doc.add_paragraph("\n.............................................................")
            doc.add_paragraph(".............................................................")

    doc.save(os.path.join(folder_path, f"de_{exam_code}.docx"))


# kết thúc hàm mới thêm vào
def export_exam_package(exam_set_id):
    # from models.exam_model import get_exam_codes_by_set, get_exam_questions_by_code, get_original_questions_by_exam_set, get_original_questions_thpt2025, get_original_questions_thpt2025, is_thpt2025_exam_set, get_exam_mcq_questions_by_code, get_exam_code_id_by_code, get_questions_by_exam_set
    

    folder_path = create_export_folder(exam_set_id)
    answer_sheet_data = {}

    exam_codes = get_exam_codes_by_set(exam_set_id)

    if is_thpt2025_exam_set(exam_set_id):
        answer_sheet_data = {}

        if exam_codes:
            for code in exam_codes:
                print("in ra code ở mỗi vòng lặp", code)
                exam_code_id = get_exam_code_id_by_code(code)
                print("in ra id của mỗi code", exam_code_id)
                questions = get_exam_questions_by_code(exam_code_id)
                print("lấy ra danh sách câu hỏi questions", questions)
                export_exam_thpt2025_to_docx(code, questions, folder_path)

                # Đáp án chỉ áp dụng cho MCQ trong phần THPT 2025
                #mcq_questions = [q for q in questions if q['type'] == 'mcq']
                mcq_questions = questions.get('mcq', [])
                answers = []
                for q in mcq_questions:
                    correct = q["correct_answer"]
                    if correct == q["A"]:
                        answers.append("A")
                    elif correct == q["B"]:
                        answers.append("B")
                    elif correct == q["C"]:
                        answers.append("C")
                    elif correct == q["D"]:
                        answers.append("D")
                    else:
                        answers.append("")
                answer_sheet_data[code] = answers

            export_answer_sheet_thpt2025_excel(answer_sheet_data, folder_path)
        else:
            original_questions = get_original_questions_thpt2025(exam_set_id)
            export_original_thpt2025_to_docx(original_questions, folder_path)
    else:
        answer_sheet_data = {}

        if exam_codes:
            for code in exam_codes:
                questions = get_exam_mcq_questions_by_code(code)
                export_exam_to_docx(code, questions, folder_path)

                answers = []
                for q in questions:
                    correct = q["correct_answer"]
                    if correct == q["answer_a"]:
                        answers.append("A")
                    elif correct == q["answer_b"]:
                        answers.append("B")
                    elif correct == q["answer_c"]:
                        answers.append("C")
                    elif correct == q["answer_d"]:
                        answers.append("D")
                    else:
                        answers.append("")

                answer_sheet_data[code] = answers

            export_answer_sheet_excel(answer_sheet_data, folder_path)
        else:
            questions = get_original_questions_by_exam_set(exam_set_id)
            export_original_questions_to_docx(questions, folder_path)

    # Nén thư mục
    zip_path = f"{folder_path}.zip"
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for root, _, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, folder_path)
                zipf.write(file_path, arcname)

    return zip_path


# CÁC HÀM MỚI THÊM VÀO ĐỂ XUẤT CÂU HỎI GỐC THÀNH FILE XLXS ĐẦU VÀO


# Hàm kiểm tra MCQ nâng cao mà chúng ta đã bàn luận
def has_advanced_mcq_features(questions: list) -> bool:
    """
    Kiểm tra xem một danh sách câu hỏi MCQ có bất kỳ câu hỏi nào
    có trường 'difficulty' (hoặc 'explanation', v.v.) không phải NULL hay không.

    Args:
        questions (list): Danh sách các dictionary câu hỏi, lấy từ bảng 'questions'.

    Returns:
        bool: True nếu có ít nhất một câu hỏi nâng cao, False nếu là MCQ thuần.
    """
    for q in questions:
        # Nếu có trường 'difficulty' và nó không phải None (tức là có giá trị)
        # hoặc nếu có trường 'explanation' và nó không rỗng
        if q.get('difficulty') is not None or (q.get('explanation') is not None and q.get('explanation').strip() != ''):
            return True
    return False

# ---

def _configure_worksheet_headers(ws, headers: list):
    """Cấu hình tiêu đề và độ rộng cột cho worksheet."""
    ws.append(headers)
    for col_num, header in enumerate(headers, 1):
        col_letter = get_column_letter(col_num)
        ws[f"{col_letter}1"].font = Font(bold=True)
        ws[f"{col_letter}1"].alignment = Alignment(horizontal='center', vertical='center')
        ws.column_dimensions[col_letter].width = 25 # Chiều rộng mặc định

def _adjust_column_widths(ws):
    """Tự động điều chỉnh độ rộng cột."""
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if cell.value is not None:
                    cell_len = len(str(cell.value))
                    if cell_len > max_length:
                        max_length = cell_len
            except:
                pass
        adjusted_width = (max_length + 2)
        if adjusted_width > 75: # Giới hạn độ rộng tối đa
            adjusted_width = 75
        elif adjusted_width < 10: # Giới hạn độ rộng tối thiểu
            adjusted_width = 10
        ws.column_dimensions[column].width = adjusted_width

# --- Các hàm xuất cụ thể ---

def export_pure_mcq_questions_to_xlsx(questions: list, folder_path: str, exam_set_id: int) -> str:
    """
    Xuất các câu hỏi MCQ thuần ra file XLSX.
    Chỉ bao gồm các cột cơ bản: Loại, Nội dung, Đáp án A-D, Đáp án đúng.

    Args:
        questions (list): Danh sách các dictionary câu hỏi từ bảng 'questions'.
        folder_path (str): Đường dẫn thư mục để lưu file XLSX.
        exam_set_id (int): ID của bộ đề, dùng để đặt tên file.

    Returns:
        str: Đường dẫn tuyệt đối đến file XLSX đã xuất.
    """
    file_name = f"BoDeThi_MCQ_Thuan_{exam_set_id}.xlsx"
    file_path = os.path.join(folder_path, file_name)

    wb = Workbook()
    ws = wb.active
    ws.title = f"MCQ Thuan - Bo De {exam_set_id}"

    headers = [
        "Loại câu hỏi", "Nội dung câu hỏi",
        "Đáp án A", "Đáp án B", "Đáp án C", "Đáp án D",
        "Đáp án đúng"
    ]
    _configure_worksheet_headers(ws, headers)

    for q in questions:
        row_data = [
            "MCQ", # Loại câu hỏi luôn là MCQ
            q.get('question_text', ''),
            q.get('answer_a', ''),
            q.get('answer_b', ''),
            q.get('answer_c', ''),
            q.get('answer_d', ''),
            q.get('correct_answer', '')
        ]
        ws.append(row_data)

    _adjust_column_widths(ws)
    wb.save(file_path)
    return file_path

def export_advanced_mcq_questions_to_xlsx(questions: list, folder_path: str, exam_set_id: int) -> str:
    """
    Xuất các câu hỏi MCQ nâng cao ra file XLSX.
    Bao gồm các cột cơ bản và thêm cột 'Giải thích', 'Mức độ khó'.

    Args:
        questions (list): Danh sách các dictionary câu hỏi từ bảng 'questions'.
        folder_path (str): Đường dẫn thư mục để lưu file XLSX.
        exam_set_id (int): ID của bộ đề, dùng để đặt tên file.

    Returns:
        str: Đường dẫn tuyệt đối đến file XLSX đã xuất.
    """
    file_name = f"BoDeThi_MCQ_NangCao_{exam_set_id}.xlsx"
    file_path = os.path.join(folder_path, file_name)

    wb = Workbook()
    ws = wb.active
    ws.title = f"MCQ Nang Cao - Bo De {exam_set_id}"

    headers = [
        "Loại câu hỏi", "Nội dung câu hỏi",
        "Đáp án A", "Đáp án B", "Đáp án C", "Đáp án D",
        "Đáp án đúng",  "Mức độ khó", "Chương" # Thêm các cột nâng cao
    ]
    _configure_worksheet_headers(ws, headers)

    for q in questions:
        row_data = [
            "MCQ", # Loại câu hỏi luôn là MCQ
            q.get('question_text', ''),
            q.get('answer_a', ''),
            q.get('answer_b', ''),
            q.get('answer_c', ''),
            q.get('answer_d', ''),
            q.get('correct_answer', ''),
            q.get('chapters', ''), # Lấy giá trị giải thích
            q.get('difficulty', '')    # Lấy giá trị độ khó
        ]
        ws.append(row_data)

    _adjust_column_widths(ws)
    wb.save(file_path)
    return file_path

# --- Hàm điều phối chính (sẽ được gọi từ route) ---

def export_exam_set_to_xlsx(user_id: int, exam_set_id: int, base_export_folder: str) -> str:
    """
    Xác định loại bộ đề và xuất câu hỏi ra file XLSX tương ứng.

    Args:
        user_id (int): ID của người dùng (để truy vấn câu hỏi).
        exam_set_id (int): ID của bộ đề cần xuất.
        base_export_folder (str): Thư mục gốc để tạo thư mục tạm thời và lưu file.

    Returns:
        str: Đường dẫn tuyệt đối đến file XLSX đã xuất.

    Raises:
        ValueError: Nếu không tìm thấy câu hỏi cho bộ đề.
        Exception: Các lỗi khác trong quá trình xuất.
    """
    # Tạo thư mục tạm thời an toàn cho mỗi lần xuất
    temp_dir = tempfile.mkdtemp(dir=base_export_folder)
    try:
        # 1. Lấy tất cả câu hỏi MCQ từ bảng 'questions'
        # Đảm bảo hàm get_questions_by_exam_set trả về danh sách dictionary
        questions_from_main_table = get_questions_by_exam_set(user_id, exam_set_id)

        if not questions_from_main_table:
            raise ValueError(f"Không tìm thấy câu hỏi cho bộ đề có ID {exam_set_id}.")

        # 2. Kiểm tra loại bộ đề
        # Phần kiểm tra THPT 2025 sẽ được thêm vào sau
        # if is_thpt2025_exam_set(exam_set_id):
        #    print(f"Bộ đề {exam_set_id} là dạng THPT 2025.")
        #    # Lấy thêm dữ liệu từ các bảng TF, SA, Essay nếu cần
        #    exported_file_path = export_thpt2025_questions_to_xlsx(exam_set_id, all_questions_data, temp_dir)
        # else:
        # Nếu không phải THPT 2025, thì nó là MCQ thuần hoặc MCQ nâng cao
        if has_advanced_mcq_features(questions_from_main_table):
            print(f"Bộ đề {exam_set_id} là dạng MCQ nâng cao.")
            exported_file_path = export_advanced_mcq_questions_to_xlsx(questions_from_main_table, temp_dir, exam_set_id)
        else:
            print(f"Bộ đề {exam_set_id} là dạng MCQ thuần.")
            exported_file_path = export_pure_mcq_questions_to_xlsx(questions_from_main_table, temp_dir, exam_set_id)

        return exported_file_path
    except Exception as e:
        # Xóa thư mục tạm thời nếu có lỗi xảy ra
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        raise e # Ném lại lỗi để route có thể bắt và xử lý
    




# KẾT THÚC HÀM XUẤT CÁC CÂU HỎI GỐC THÀNH FIEL XLXS


# ĐÂY LÀ HÀM IMPORT CÂU HỎI TỪ FILE XLXS


def import_questions_from_xlsx(file_path: str, exam_set_id: int, user_id: int) -> dict:
    """
    Nhập câu hỏi MCQ từ file XLSX vào cơ sở dữ liệu.

    Args:
        file_path (str): Đường dẫn đến file XLSX cần nhập.
        exam_set_id (int): ID của bộ đề để thêm câu hỏi vào (đã tạo trước).
        user_id (int): ID của người dùng tạo câu hỏi.

    Returns:
        dict: Kết quả của quá trình nhập (số lượng câu hỏi đã nhập, lỗi...).
    """
    
    results = {
        "total_rows": 0,
        "imported_questions": 0,
        "failed_rows": [],
        "messages": []
    }

    try:
        wb = openpyxl.load_workbook(file_path)
        ws = wb.active

        # Tạo map từ tên cột sang chỉ số cột (0-based)
        col_map = {cell.value.strip(): idx for idx, cell in enumerate(ws[1]) if cell.value}

        # Các cột bắt buộc
        required_cols = ["Nội dung câu hỏi", "Đáp án A", "Đáp án B", "Đáp án C", "Đáp án D", "Đáp án đúng"]
        if not all(col in col_map for col in required_cols):
            results["messages"].append("❌ File thiếu các cột bắt buộc: " + ", ".join(required_cols))
            return results

        has_difficulty_col = "Mức độ khó" in col_map
        has_chapter_col = "Chapter" in col_map

        # Duyệt từng hàng dữ liệu (bỏ qua hàng đầu)
        for row_idx in range(2, ws.max_row + 1):
            results["total_rows"] += 1
            row_values = [cell.value for cell in ws[row_idx]]

            # Lấy dữ liệu từng cột
            def get(col_name):
                idx = col_map.get(col_name)
                if idx is not None and idx < len(row_values):
                    val = row_values[idx]
                    return str(val).strip() if val is not None else ''
                return ''

            question_text = get("Nội dung câu hỏi")
            answer_a = get("Đáp án A")
            answer_b = get("Đáp án B")
            answer_c = get("Đáp án C")
            answer_d = get("Đáp án D")
            correct_answer = get("Đáp án đúng")

            # Tùy chọn
            difficulty = get("Mức độ khó") if has_difficulty_col else None
            chapter = get("Chapter") if has_chapter_col else None

            # Kiểm tra dữ liệu bắt buộc
            if not question_text or not correct_answer or not (answer_a or answer_b or answer_c or answer_d):
                results["failed_rows"].append(
                    f"Hàng {row_idx}: Thiếu nội dung câu hỏi, đáp án đúng, hoặc ít nhất một đáp án A–D."
                )
                continue

            # Chuẩn bị dữ liệu câu hỏi
            question_data = {
                'exam_set_id': exam_set_id,
                'user_id': user_id,
                'question_text': question_text,
                'answer_a': answer_a,
                'answer_b': answer_b,
                'answer_c': answer_c,
                'answer_d': answer_d,
                'correct_answer': correct_answer,
                'difficulty': difficulty if difficulty else None,
                'chapter': chapter if chapter else None
            }

            try:
                insert_mcq_question(question_data)
                results["imported_questions"] += 1

            except ValueError as ve:
                results["failed_rows"].append(f"Hàng {row_idx}: Lỗi dữ liệu – {ve}")
                results["messages"].append(f"Lỗi hàng {row_idx}: {ve}")

            except Exception as e:
                results["failed_rows"].append(f"Hàng {row_idx}: Lỗi khi lưu vào DB – {e}")
                results["messages"].append(f"Lỗi hàng {row_idx}: {e}")

    except Exception as e:
        results["messages"].append(f"⛔️ Lỗi khi mở hoặc xử lý file Excel: {e}")

    finally:
        if os.path.exists(file_path):
            os.remove(file_path)

    return results
    

# KẾT THÚC HÀM IMPORT CÂU HỎI TỪ FILE XLXS

def render_exam_to_docx(data):
    
    doc = Document()

# Thiết lập font mặc định: Times New Roman, 13pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# ==== THÔNG TIN ĐẦU TRANG ==== 

    p = doc.add_paragraph("TRƯỜNG ĐẠI HỌC KIÊN GIANG")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.runs[0].bold = True

    p = doc.add_paragraph("[ĐƠN VỊ QUẢN LÝ HỌC PHẦN]")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.runs[0].italic = True

    p = doc.add_paragraph("ĐỀ THI KẾT THÚC HỌC PHẦN")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("Học kỳ: ……  Năm học: ……… - ………")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # dòng trống

    doc.add_paragraph(f"ĐỀ THI - MÃ ĐỀ {data['exam_code']}", 0)
    doc.add_paragraph("Tên học phần: ..................................................        Mã học phần: ...........................")

# Hình thức thi + thời gian
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Hình thức thi: ")
    run.bold = True
    p.add_run("Trắc nghiệm khách quan")

    p = doc.add_paragraph()
    p.add_run("Thời gian làm bài: ").bold = True
    p.add_run("……… phút (không kể thời gian phát đề)")

# Ghi chú
    p = doc.add_paragraph("Ghi chú:")
    p.add_run("\n    - Thí sinh được/không được sử dụng tài liệu.")
    p.add_run("\n    - Thí sinh nộp lại đề thi.")

# Tiêu đề nội dung
    p = doc.add_paragraph("\nNỘI DUNG ĐỀ THI")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].bold = True
    
    
    ## kết thúc thêm vào
    #doc.add_heading(f"ĐỀ THI - MÃ ĐỀ {data['exam_code']}", 0)
    
    
    questions = []

    for q in data.get('mcq_questions', []):
        q['type'] = 'MCQ'
        questions.append(q)

    for q in data.get('tf_questions', []):
        q['type'] = 'TF'
        questions.append(q)

    for q in data.get('sa_questions', []):
        q['type'] = 'SA'
        questions.append(q)

# (Tùy chọn) Sắp xếp lại theo thứ tự nếu có 'position'
    questions.sort(key=lambda q: q.get('position', 0))

# Cập nhật lại vào data để truyền vào hàm render
    data['questions'] = questions

# Gọi render như bình thường
    

    for idx, q in enumerate(questions, start=1):
        doc.add_paragraph(f"Câu {idx}: {q['question_text']}", style='Normal')
        
        if q['type'] == 'MCQ':
            doc.add_paragraph(f"A. {q['answer_a']}")
            doc.add_paragraph(f"B. {q['answer_b']}")
            doc.add_paragraph(f"C. {q['answer_c']}")
            doc.add_paragraph(f"D. {q['answer_d']}")

        elif q['type'] == 'TF':
            doc.add_paragraph("Các phát biểu:")
            for i, st in enumerate(q['statements'], start=1):
                doc.add_paragraph(f"  ({i}) {st['label']}: {st['content']}")

        elif q['type'] == 'SA':
            doc.add_paragraph("(Câu hỏi tự luận. Học sinh trả lời ngắn vào chỗ trống.)")

        doc.add_paragraph("")

    return doc

def render_exam_to_pdf(data):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    margin_left = 30
    line_height = 16
    y = height - 40

    def draw_text(text, size=12, x=None, align='left'):
        c.setFont('HeiseiMin-W3', size)  # Dùng font Unicode
        if align == 'center':
            c.drawCentredString(width / 2, y, text)
        elif align == 'right':
            c.drawRightString(width - margin_left, y, text)
        else:
            c.drawString(margin_left if x is None else x, y, text)

    # ==== HEADER ====
    draw_text("TRƯỜNG ĐẠI HỌC KIÊN GIANG", size=13); y -= line_height
    draw_text("[ĐƠN VỊ QUẢN LÝ HỌC PHẦN]", size=13); y -= line_height
    draw_text("ĐỀ THI KẾT THÚC HỌC PHẦN", size=13, align='center'); y -= line_height
    draw_text("Học kỳ: ……     Năm học: ……… - ………", size=12, align='center'); y -= 2 * line_height
    draw_text(f"Mã đề: {data.get('code', '……')}", size=12); y -= line_height
    draw_text("Tên học phần: .........................................     Mã học phần: .......................", size=12); y -= line_height
    draw_text("Hình thức thi: Trắc nghiệm khách quan", size=12); y -= line_height
    draw_text("Thời gian làm bài: …… phút (không kể thời gian phát đề)", size=12); y -= 2 * line_height
    draw_text("Ghi chú:", size=12); y -= line_height
    draw_text("- Thí sinh được/không được sử dụng tài liệu.", size=12); y -= line_height
    draw_text("- Thí sinh nộp lại đề thi.", size=12); y -= 2 * line_height
    draw_text("NỘI DUNG ĐỀ THI", size=13, align='center'); y -= 2 * line_height

    # Gom các câu hỏi
    questions = []
    for q in data.get('mcq_questions', []):
        q['type'] = 'MCQ'
        questions.append(q)
    for q in data.get('tf_questions', []):
        q['type'] = 'TF'
        questions.append(q)
    for q in data.get('sa_questions', []):
        q['type'] = 'SA'
        questions.append(q)
    questions.sort(key=lambda q: q.get('position', 0))
    data['questions'] = questions

    # In nội dung câu hỏi
    for idx, q in enumerate(data['questions'], start=1):
        if y < 100:
            c.showPage()
            y = height - 50
            c.setFont('HeiseiMin-W3', 12)

        c.drawString(50, y, f"Câu {idx}: {q['question_text']}")
        y -= 20

        if q['type'] == 'MCQ':
            c.drawString(60, y, f"A. {q['answer_a']}"); y -= 15
            c.drawString(60, y, f"B. {q['answer_b']}"); y -= 15
            c.drawString(60, y, f"C. {q['answer_c']}"); y -= 15
            c.drawString(60, y, f"D. {q['answer_d']}"); y -= 20

        elif q['type'] == 'TF':
            c.drawString(60, y, "Các phát biểu:"); y -= 15
            for i, st in enumerate(q['statements'], start=1):
                c.drawString(70, y, f"({i}) {st['label']}: {st['content']}")
                y -= 15
            y -= 5

        elif q['type'] == 'SA':
            c.drawString(60, y, "(Câu hỏi tự luận - học sinh trả lời vào giấy.)")
            y -= 20

    c.save()
    buffer.seek(0)
    return buffer

def generate_score_excel(room_code, room_info, score_data):
    wb = Workbook()
    ws = wb.active
    ws.title = f"Bảng điểm {room_code}"

    # Thêm thông tin phòng thi lên đầu
    ws.append(["TRƯỜNG ĐẠI HỌC .........................................."])
    ws.append(["KHOA: .................................................."])
    ws.append(["TP. ....................................................."])
    ws.append([])

    ws.append(["BẢNG ĐIỂM THI KẾT THÚC HỌC PHẦN"])
    ws.append([f"Môn: {room_info['subject_name']}"])
    ws.append([f"Mã phòng: {room_info['room_code']}"])
    ws.append([f"Ngày thi: {room_info['open_time'].strftime('%d/%m/%Y') if room_info['open_time'] else ''}"])
    ws.append([f"Thời gian làm bài: {room_info['duration_minutes']} phút"])
    ws.append([f"Giảng viên: {room_info['instructor_name'] or '........................'}"])
    ws.append([])

    # Header bảng điểm
    headers = ["STT", "Mã sinh viên", "Họ tên", "Mã đề", "Điểm", "Bắt đầu lúc", "Kết thúc lúc", "Trạng thái"]
    ws.append(headers)

    # Dữ liệu điểm
    for idx, row in enumerate(score_data, start=1):
        ws.append([
            idx,
            row['student_code'],
            row['student_name'] or row.get('full_name') or 'Không rõ',
            row['exam_code'] or '',
            float(row['score']) if row['score'] is not None else '',
            row['start_time'].strftime('%H:%M %d/%m/%Y') if row['start_time'] else '',
            row['end_time'].strftime('%H:%M %d/%m/%Y') if row['end_time'] else '',
            row['status'].capitalize() if row['status'] else ''
        ])

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    filename = f"bang_diem_{room_code}.xlsx"
    return send_file(output, as_attachment=True, download_name=filename, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")



## MỚI THÊM VÀO

def render_answer_to_docx(data):
    doc = Document()

    # Font mặc định Times New Roman 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Header
    doc.add_paragraph("TRƯỜNG ĐẠI HỌC KIÊN GIANG", style='Normal').bold = True
    doc.add_paragraph("[ĐƠN VỊ QUẢN LÝ HỌC PHẦN]", style='Normal')
    doc.add_paragraph("ĐÁP ÁN ĐỀ THI KẾT THÚC HỌC PHẦN TRẮC NGHIỆM KHÁCH QUAN", style='Normal').bold = True
    doc.add_paragraph("Học kỳ: …… Năm học: ……… - ………")
    doc.add_paragraph("Tên học phần: ...................................................  Mã học phần: ...................  Số TC: ........")
    doc.add_paragraph("Ngành: .................................................  Hệ đào tạo: ............  Hình thức đào tạo: ...............")
    doc.add_paragraph("")

    # Bảng đáp án
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Câu hỏi'
    hdr_cells[1].text = 'Đáp án đúng'
    hdr_cells[2].text = 'Loại'

    questions = []

    for q in data.get('mcq_questions', []):
        q['type'] = 'MCQ'
        questions.append(q)
    for q in data.get('tf_questions', []):
        q['type'] = 'TF'
        questions.append(q)
    for q in data.get('sa_questions', []):
        q['type'] = 'SA'
        questions.append(q)

    questions.sort(key=lambda q: q.get('position', 0))

    for idx, q in enumerate(questions, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)

        if q['type'] == 'MCQ':
            row_cells[1].text = q.get('correct_answer', '')
            row_cells[2].text = 'MCQ'
        elif q['type'] == 'TF':
            corrects = [s['label'] for s in q['statements'] if s['is_true']]
            row_cells[1].text = ", ".join(corrects)
            row_cells[2].text = 'TF'
        elif q['type'] == 'SA':
            row_cells[1].text = q.get('answer_text', '')
            row_cells[2].text = 'SA'

    doc.add_paragraph("HẾT")
    return doc


def render_exam_to_pdf(data):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    margin_left = 30
    line_height = 16
    y = height - 40

    def draw_text(text, size=12, x=None, align='left'):
        c.setFont('Roboto', size)
        if align == 'center':
            c.drawCentredString(width / 2, y, text)
        elif align == 'right':
            c.drawRightString(width - margin_left, y, text)
        else:
            c.drawString(margin_left if x is None else x, y, text)

    # ==== HEADER MẪU ====
    draw_text("TRƯỜNG ĐẠI HỌC KIÊN GIANG", size=13); y -= line_height
    draw_text("[ĐƠN VỊ QUẢN LÝ HỌC PHẦN]", size=13); y -= line_height
    draw_text("ĐỀ THI KẾT THÚC HỌC PHẦN", size=13, align='center'); y -= line_height
    draw_text("Học kỳ: ……     Năm học: ……… - ………", size=12, align='center'); y -= 2 * line_height
    draw_text(f"Mã đề: {data.get('code', '……')}", size=12); y -= line_height
    draw_text("Tên học phần: .........................................     Mã học phần: .......................", size=12); y -= line_height
    draw_text("Hình thức thi: Trắc nghiệm khách quan", size=12); y -= line_height
    draw_text("Thời gian làm bài: …… phút (không kể thời gian phát đề)", size=12); y -= 2 * line_height
    draw_text("Ghi chú:", size=12); y -= line_height
    draw_text("- Thí sinh được/không được sử dụng tài liệu.", size=12); y -= line_height
    draw_text("- Thí sinh nộp lại đề thi.", size=12); y -= 2 * line_height
    draw_text("NỘI DUNG ĐỀ THI", size=13, align='center'); y -= 2 * line_height

    # ==== CÂU HỎI ====
    questions = []
    for q in data.get('mcq_questions', []):
        q['type'] = 'MCQ'
        questions.append(q)
    for q in data.get('tf_questions', []):
        q['type'] = 'TF'
        questions.append(q)
    for q in data.get('sa_questions', []):
        q['type'] = 'SA'
        questions.append(q)

    # Sắp xếp nếu có thứ tự
    questions.sort(key=lambda q: q.get('position', 0))
    data['questions'] = questions

    for idx, q in enumerate(data['questions'], start=1):
        if y < 100:
            c.showPage()
            y = height - 50
            c.setFont('Roboto', 12)

        c.drawString(50, y, f"Câu {idx}: {q['question_text']}")
        y -= 20

        if q['type'] == 'MCQ':
            c.drawString(60, y, f"A. {q['answer_a']}"); y -= 15
            c.drawString(60, y, f"B. {q['answer_b']}"); y -= 15
            c.drawString(60, y, f"C. {q['answer_c']}"); y -= 15
            c.drawString(60, y, f"D. {q['answer_d']}"); y -= 20

        elif q['type'] == 'TF':
            c.drawString(60, y, "Các phát biểu:"); y -= 15
            for i, st in enumerate(q['statements'], start=1):
                c.drawString(70, y, f"({i}) {st['label']}: {st['content']}")
                y -= 15
            y -= 5

        elif q['type'] == 'SA':
            c.drawString(60, y, "(Câu hỏi tự luận - học sinh trả lời vào giấy.)")
            y -= 20

    c.save()
    buffer.seek(0)
    return buffer

### mới thêm vào đẻ xuất cả câu hỏi lẫn đáp án



def render_exam_with_answers_to_docx(data):
    doc = Document()

    # ==== FONT & STYLE CHUNG ====
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # ==== PHẦN ĐỀ THI ====
    doc.add_paragraph("TRƯỜNG ĐẠI HỌC KIÊN GIANG").runs[0].bold = True
    doc.add_paragraph("[ĐƠN VỊ QUẢN LÝ HỌC PHẦN]").runs[0].italic = True
    doc.add_paragraph("ĐỀ THI KẾT THÚC HỌC PHẦN").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("Học kỳ: ……  Năm học: ……… - ………")
    doc.add_paragraph(f"\nĐỀ THI - MÃ ĐỀ {data['exam_code']}")
    doc.add_paragraph("Tên học phần: ..................................................        Mã học phần: ...........................")
    doc.add_paragraph("Thời gian làm bài: ……… phút (không kể thời gian phát đề)")
    doc.add_paragraph("Hình thức thi: Trắc nghiệm khách quan")
    doc.add_paragraph("Ghi chú:\n    - Thí sinh được/không được sử dụng tài liệu.\n    - Thí sinh nộp lại đề thi.")
    doc.add_paragraph("\nNỘI DUNG ĐỀ THI").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ==== PHẦN CÂU HỎI ====
    questions = []
    for q in data.get('mcq_questions', []):
        q['type'] = 'MCQ'
        questions.append(q)
    for q in data.get('tf_questions', []):
        q['type'] = 'TF'
        questions.append(q)
    for q in data.get('sa_questions', []):
        q['type'] = 'SA'
        questions.append(q)
    
    questions.sort(key=lambda q: q.get('position', 0))
    data['questions'] = questions

    for idx, q in enumerate(questions, start=1):
        doc.add_paragraph(f"Câu {idx}: {q['question_text']}")
        if q['type'] == 'MCQ':
            doc.add_paragraph(f"A. {q['answer_a']}")
            doc.add_paragraph(f"B. {q['answer_b']}")
            doc.add_paragraph(f"C. {q['answer_c']}")
            doc.add_paragraph(f"D. {q['answer_d']}")
        elif q['type'] == 'TF':
            image_url = q.get('image_url')
            if image_url:
                try:
                    relative_path = image_url.lstrip("/")
                    if relative_path.startswith("uploads/"):
                        relative_path = relative_path.replace("uploads/", "")
                    image_path = os.path.join(UPLOAD_FOLDER, relative_path)
                    doc.add_picture(image_path, width=Inches(4.5))
                except Exception as e:
                    doc.add_paragraph(f"(Không thể tải ảnh: {str(e)})")
            doc.add_paragraph("Các phát biểu:")
            for i, st in enumerate(q['statements'], start=1):
                doc.add_paragraph(f"  ({i}) {st['label']}: {st['content']}")
        elif q['type'] == 'SA':
            doc.add_paragraph("(Câu hỏi tự luận. Học sinh trả lời ngắn vào chỗ trống.)")
        doc.add_paragraph("")  # dòng trống giữa câu hỏi

    # ==== PHẦN ĐÁP ÁN ====
    doc.add_paragraph("\nĐÁP ÁN").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Câu hỏi"
    hdr_cells[1].text = "Đáp án đúng"
    

    for idx, q in enumerate(questions, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        if q['type'] == 'MCQ':
            row_cells[1].text = q.get('correct_answer', '')
            
        elif q['type'] == 'TF':
            answer_lines = []
            for s in q['statements']:
                label = s.get("label", "?")
                truth = "Đúng" if s.get("is_true") else "Sai"
                answer_lines.append(f"{label}. {truth}")
            row_cells[1].text = "\n".join(answer_lines)
           
        elif q['type'] == 'SA':
            row_cells[1].text = q.get('correct_answer', '')
            

    doc.add_paragraph("HẾT")

    return doc

### HÀM NÀY LÀ ĐỂ XUẤT PDF ĐÁP ÁN CÂU HỎI GỐC


def export_exam_with_answers_to_pdf(data):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    flow = []

    # ===== HEADER =====
    flow.append(Paragraph("TRƯỜNG ĐẠI HỌC KIÊN GIANG", styles['RobotoTitle']))
    flow.append(Paragraph("[ĐƠN VỊ QUẢN LÝ HỌC PHẦN]", styles['RobotoNormal']))
    flow.append(Paragraph("ĐỀ THI KẾT THÚC HỌC PHẦN", styles['RobotoTitle']))
    flow.append(Spacer(1, 12))
    flow.append(Paragraph(f"ĐỀ THI - MÃ ĐỀ {data.get('exam_code', '')}", styles['RobotoHeading2']))
    flow.append(Spacer(1, 12))
    flow.append(Paragraph("NỘI DUNG ĐỀ THI", styles['RobotoHeading2']))
    flow.append(Spacer(1, 12))
    
    
    questions = []
    for q in data.get('mcq_questions', []):
        q['type'] = 'MCQ'
        questions.append(q)
    for q in data.get('tf_questions', []):
        q['type'] = 'TF'
        questions.append(q)
    for q in data.get('sa_questions', []):
        q['type'] = 'SA'
        questions.append(q)

    # Sắp xếp nếu có thứ tự
    questions.sort(key=lambda q: q.get('position', 0))
    data['questions'] = questions
    
    
    
    

    questions = data.get("questions", [])
    for idx, q in enumerate(questions, start=1):
        flow.append(Paragraph(f"Câu {idx}: {q.get('question_text', '')}", styles['RobotoNormal']))

        if q['type'] == 'MCQ':
            flow.append(Paragraph(f"A. {q.get('answer_a', '')}", styles['RobotoNormal']))
            flow.append(Paragraph(f"B. {q.get('answer_b', '')}", styles['RobotoNormal']))
            flow.append(Paragraph(f"C. {q.get('answer_c', '')}", styles['RobotoNormal']))
            flow.append(Paragraph(f"D. {q.get('answer_d', '')}", styles['RobotoNormal']))

        elif q['type'] == 'TF':
            image_url = q.get('image_url')
            if image_url:
                try:
                    relative_path = image_url.lstrip("/")
                    if relative_path.startswith("uploads/"):
                        relative_path = relative_path.replace("uploads/", "")
                    image_path = os.path.join(UPLOAD_FOLDER, relative_path)

                    image = Image(image_path, width=300)
                    flow.append(image)
                    flow.append(Spacer(1, 6))
                except Exception as e:
                    flow.append(Paragraph(f"(Không thể tải ảnh: {str(e)})", styles['RobotoNormal']))

            flow.append(Paragraph("Các phát biểu:", styles['RobotoNormal']))
            for i, st in enumerate(q.get('statements', []), start=1):
                label = st.get("label", "?")
                content = st.get("content", "")
                flow.append(Paragraph(f"({i}) {label}: {content}", styles['RobotoNormal']))

        elif q['type'] == 'SA':
            flow.append(Paragraph("(Câu hỏi tự luận. Học sinh trả lời ngắn vào chỗ trống.)", styles['RobotoNormal']))

        flow.append(Spacer(1, 12))

    # ===== BẢNG ĐÁP ÁN =====
    flow.append(Paragraph("BẢNG ĐÁP ÁN", styles['RobotoHeading2']))
    flow.append(Spacer(1, 12))

    table_data = [['Câu', 'Đáp án đúng']]
    for idx, q in enumerate(questions, start=1):
        if q['type'] == 'MCQ':
            answer = q.get('correct_answer', '')
        elif q['type'] == 'TF':
            answer_lines = []
            for s in q.get('statements', []):
                label = s.get("label", "?")
                truth = "Đúng" if s.get("is_true") else "Sai"
                answer_lines.append(f"{label}. {truth}")
            answer = "\n".join(answer_lines)
        elif q['type'] == 'SA':
            answer = q.get('answer_text', '')
        else:
            answer = ''
        table_data.append([str(idx), answer])

    table = Table(table_data, colWidths=[50, 400])
    table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Roboto'),
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
    ]))
    flow.append(table)
    flow.append(Spacer(1, 24))
    flow.append(Paragraph("HẾT", styles['RobotoTitle']))

    doc.build(flow)
    buffer.seek(0)
    return buffer

### đây là xuất docx câu hỏi và đáp án



def render_answer_to_docx_advance(data):
    doc = Document()

    # ==== Font mặc định Times New Roman 12pt ====
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # ==== Header ====
    p = doc.add_paragraph("TRƯỜNG ĐẠI HỌC KIÊN GIANG")
    p.runs[0].bold = True

    doc.add_paragraph("[ĐƠN VỊ QUẢN LÝ HỌC PHẦN]")
    p = doc.add_paragraph("ĐÁP ÁN ĐỀ THI KẾT THÚC HỌC PHẦN TRẮC NGHIỆM KHÁCH QUAN")
    p.runs[0].bold = True

    doc.add_paragraph("Học kỳ: ……  Năm học: ……… - ………")
    doc.add_paragraph("Tên học phần: ...................................................  Mã học phần: ...................  Số TC: ........")
    doc.add_paragraph("Ngành: .................................................  Hệ đào tạo: ............  Hình thức đào tạo: ...............")
    doc.add_paragraph("")  # dòng trống

    # ==== Tổng hợp dữ liệu câu hỏi ====
    questions = []

    for q in data.get('mcq_questions', []):
        q['type'] = 'MCQ'
        questions.append(q)

    for q in data.get('tf_questions', []):
        q['type'] = 'TF'
        questions.append(q)

    for q in data.get('sa_questions', []):
        q['type'] = 'SA'
        questions.append(q)

    questions.sort(key=lambda q: q.get('position', 0))

    # ==== BẢNG ĐÁP ÁN ====
    doc.add_paragraph("BẢNG ĐÁP ÁN").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Câu hỏi"
    hdr_cells[1].text = "Đáp án đúng"

    for idx, q in enumerate(questions, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)

        if q['type'] == 'MCQ':
            row_cells[1].text = q.get('correct_answer', '')
        elif q['type'] == 'TF':
            
             image_url = q.get('image_url')
             if image_url:
                try:
                # Build full URL nếu cần
                    if image_url.startswith("/"):
                        image_url = request.host_url.rstrip("/") + image_url

                    response = requests.get(image_url)
                    response.raise_for_status()
                    image_stream = io.BytesIO(response.content)

                # Thêm ảnh sau dòng bảng
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(image_stream, width=Inches(4.5))
                except Exception as e:
                    doc.add_paragraph(f"(Không thể tải ảnh: {str(e)})")
                    corrects = [s['label'] for s in q['statements'] if s['is_true']]
                    row_cells[1].text = ", ".join(corrects)
        elif q['type'] == 'SA':
            row_cells[1].text = q.get('answer_text', '')

    doc.add_paragraph("")
    doc.add_paragraph("HẾT").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return doc

##### Hàm này là để xuất bài thi sinh viên ra file

def export_student_exam_to_docx(exam_info, questions):
    doc = Document()

    def add_heading(text):
        run = doc.add_paragraph().add_run(text)
        run.bold = True
        run.font.size = Pt(14)

    def add_colored_text(paragraph, text, color_rgb):
        run = paragraph.add_run(text)
        run.font.size = Pt(12)
        run.font.color.rgb = color_rgb

    # --- Header giống giấy thi ---
    doc.add_heading("PHIẾU BÀI LÀM", level=0)
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    table.autofit = True

    table.cell(0, 0).text = f"Họ tên: {exam_info['student_name']}"
    table.cell(0, 1).text = f"Mã SV: {exam_info['student_code']}"
    table.cell(0, 2).text = f"Phòng thi: {exam_info['room_code']}"
    table.cell(0, 3).text = f"Mã đề: {exam_info['exam_code']}"

    table.cell(1, 0).text = f"Môn: {exam_info['subject_name']}"
    table.cell(1, 1).text = f"Thời gian: {exam_info['duration_minutes']} phút"
    table.cell(1, 2).text = f"Bắt đầu: {exam_info['start_time']}"
    table.cell(1, 3).text = f"Kết thúc: {exam_info['end_time']}"

    doc.add_paragraph()  # dòng trắng

    # --- Hiển thị các câu hỏi ---
    for idx, q in enumerate(questions, start=1):
        doc.add_paragraph(f"Câu {idx} ({q['type']}):", style="List Number")

        if q['type'] == 'MCQ':
            doc.add_paragraph(q['question_text'])

            for opt in q['options']:
                p = doc.add_paragraph()
                label = opt['label']
                text = f"{label}. {opt['text']}"

                if q['selected_answer'] == label:
                    if q['is_correct']:
                        add_colored_text(p, text, RGBColor(0, 128, 0))  # Xanh lá
                    else:
                        add_colored_text(p, text, RGBColor(255, 0, 0))  # Đỏ
                elif q['correct_answer'] == label:
                    add_colored_text(p, text, RGBColor(0, 128, 0))  # Tô xanh đáp án đúng
                else:
                    p.add_run(text)

        elif q['type'] == 'TF':
            doc.add_paragraph(q['question_text'])

            if 'image_url' in q and q['image_url']:
                try:
                    # img_data = requests.get(q['image_url']).content
                    # doc.add_picture(io.BytesIO(img_data), width=Inches(4.5))
                    image_url = q['image_url']
                    if image_url.startswith("/"):
                        image_url = request.host_url.rstrip("/") + image_url

                    response = requests.get(image_url)
                    response.raise_for_status()
                    doc.add_picture(io.BytesIO(response.content), width=Inches(4.5))
                except:
                    doc.add_paragraph("(Không thể tải hình ảnh)")

            for stmt in q['statements']:
                content = f"{stmt['label']}. {stmt['content']}"
                answer = "Đúng" if stmt.get('student_answer') else "Sai"
                is_correct = stmt.get('is_correct')

                p = doc.add_paragraph()
                p.add_run(content + " - ")
                if is_correct is True:
                    add_colored_text(p, f"Trả lời: {answer}", RGBColor(0, 128, 0))
                elif is_correct is False:
                    add_colored_text(p, f"Trả lời: {answer}", RGBColor(255, 0, 0))
                    correct_text = "Đúng" if stmt['is_true'] else "Sai"
                    p.add_run(f" (Đáp án đúng: {correct_text})")

        elif q['type'] == 'SA':
            doc.add_paragraph(q['question_text'])
            answer = q.get('student_answer') or "(Không trả lời)"
            p = doc.add_paragraph("Trả lời: ")
            if q.get('is_correct') is True:
                add_colored_text(p, answer, RGBColor(0, 128, 0))
            elif q.get('is_correct') is False:
                add_colored_text(p, answer, RGBColor(255, 0, 0))
            else:
                p.add_run(answer)

        doc.add_paragraph()  # khoảng cách giữa các câu

    # --- Xuất file ---
    # output = io.BytesIO()
    # doc.save(output)
    # output.seek(0)

    # filename = f"BaiLam_{exam_info['student_code']}_De{exam_info['exam_code']}.docx"
    return doc


####### kết thúc hàm xuarasrt bài thi sinh viên

##### hàm này để xuất cau hỏi và đáp án dạng docx gốc

def export_original_exam_to_docx(exam_type, data):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.add_paragraph("ĐỀ THI GỐC").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")

    idx = 1

    if exam_type == 'mcq':
        for q in data:
            doc.add_paragraph(f"Câu {idx}: {q['question_text']}")
            doc.add_paragraph(f"A. {q.get('answer_a', '')}")
            doc.add_paragraph(f"B. {q.get('answer_b', '')}")
            doc.add_paragraph(f"C. {q.get('answer_c', '')}")
            doc.add_paragraph(f"D. {q.get('answer_d', '')}")
            idx += 1
            doc.add_paragraph("")

    elif exam_type == 'thpt_2025':
        # MCQ
        for q in data.get('MCQ', []):
            doc.add_paragraph(f"Câu {idx}: {q['question_text']}")
            doc.add_paragraph(f"A. {q['answer_a']}")
            doc.add_paragraph(f"B. {q['answer_b']}")
            doc.add_paragraph(f"C. {q['answer_c']}")
            doc.add_paragraph(f"D. {q['answer_d']}")
            idx += 1
            doc.add_paragraph("")

        # TF
        for q in data.get('tf', []):
            doc.add_paragraph(f"Câu {idx} (Đúng/Sai): {q['question_text']}")
            #image_url = q.get('image_url')
            if 'image_url' in q and q['image_url']:
                try:
                    relative_path = q['image_url'].lstrip("/").replace("uploads/", "")
                    image_path = os.path.join(UPLOAD_FOLDER, relative_path)
                    doc.add_picture(image_path, width=Pt(200))
                except Exception as e:
                    doc.add_paragraph("[Không thể hiển thị ảnh]")
                    print("❌ Lỗi hiển thị ảnh:", e)

            for i, st in enumerate(q['sub_items'], start=1):
                doc.add_paragraph(f"({i}) {st['statement_text']}")
            idx += 1
            doc.add_paragraph("")

        # SA
        for q in data.get('SA', []):
            doc.add_paragraph(f"Câu {idx} (Tự luận): {q['question_text']}")
            idx += 1
            doc.add_paragraph("")

    # ==== BẢNG ĐÁP ÁN ====
    doc.add_paragraph("BẢNG ĐÁP ÁN").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Câu"
    table.rows[0].cells[1].text = "Đáp án đúng"

    idx = 1

    if exam_type == 'mcq':
        for q in data:
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = q.get('correct_answer', '')
            idx += 1

    else:
        for q in data.get('MCQ', []):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = q.get('correct_answer', '')
            idx += 1

        for q in data.get('tf', []):
            corrects = [s['statement_text'] for s in q['sub_items'] if s['is_true']]
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = ", ".join(corrects)
            idx += 1

        for q in data.get('SA', []):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = q.get('correct_answer', '')
            idx += 1

    return doc

#### hàm này xuất dạng pdf



def export_original_exam_to_pdf(exam_type, data):
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    flow = []

    flow.append(Paragraph("ĐỀ THI GỐC", styles['Title']))
    flow.append(Spacer(1, 12))

    idx = 1

    if exam_type == 'mcq':
        for q in data:
            flow.append(Paragraph(f"Câu {idx}: {q['question_text']}", styles['Normal']))
            flow.append(Paragraph(f"A. {q.get('answer_a', '')}", styles['Normal']))
            flow.append(Paragraph(f"B. {q.get('answer_b', '')}", styles['Normal']))
            flow.append(Paragraph(f"C. {q.get('answer_c', '')}", styles['Normal']))
            flow.append(Paragraph(f"D. {q.get('answer_d', '')}", styles['Normal']))
            flow.append(Spacer(1, 12))
            idx += 1

    else:
        for q in data.get('MCQ', []):
            flow.append(Paragraph(f"Câu {idx}: {q['question_text']}", styles['Normal']))
            flow.append(Paragraph(f"A. {q['answer_a']}", styles['Normal']))
            flow.append(Paragraph(f"B. {q['answer_b']}", styles['Normal']))
            flow.append(Paragraph(f"C. {q['answer_c']}", styles['Normal']))
            flow.append(Paragraph(f"D. {q['answer_d']}", styles['Normal']))
            flow.append(Spacer(1, 12))
            idx += 1

        for q in data.get('tf', []):
            flow.append(Paragraph(f"Câu {idx} (Đúng/Sai): {q['question_text']}", styles['Normal']))
            if 'image_url' in q and q['image_url']:
                try:
                    flow.append(Image(q['image_url'], width=200, height=150))
                except Exception as e:
                    flow.append(Paragraph(f"[Lỗi ảnh: {str(e)}]", styles['Italic']))
            for i, st in enumerate(q['sub_items'], start=1):
                flow.append(Paragraph(f"({i}) {st['statement_text']}", styles['Normal']))
            flow.append(Spacer(1, 12))
            idx += 1

        for q in data.get('SA', []):
            flow.append(Paragraph(f"Câu {idx} (Tự luận): {q['question_text']}", styles['Normal']))
            flow.append(Spacer(1, 12))
            idx += 1

    flow.append(Paragraph("BẢNG ĐÁP ÁN", styles['Heading2']))
    flow.append(Spacer(1, 12))

    table_data = [["Câu", "Đáp án đúng"]]
    idx = 1

    if exam_type == 'mcq':
        for q in data:
            table_data.append([str(idx), q.get('correct_answer', '')])
            idx += 1
    else:
        for q in data.get('MCQ', []):
            table_data.append([str(idx), q.get('correct_answer', '')])
            idx += 1
        for q in data.get('tf', []):
            corrects = [s['statement_text'] for s in q['sub_items'] if s['is_true']]
            table_data.append([str(idx), ", ".join(corrects)])
            idx += 1
        for q in data.get('SA', []):
            table_data.append([str(idx), q.get('correct_answer', '')])
            idx += 1

    table = Table(table_data, colWidths=[50, 400])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('GRID', (0,0), (-1,-1), 0.5, colors.black),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
    ]))
    flow.append(table)
    flow.append(Spacer(1, 12))
    flow.append(Paragraph("HẾT", styles['Title']))

    doc.build(flow)
    buf.seek(0)
    return buf
def export_original_exam_to_pdf_advance(exam_type, data):
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    flow = []
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    y = height - 50

    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(width / 2, y, "TRƯỜNG ĐẠI HỌC KIÊN GIANG")
    y -= 15
    c.setFont("Helvetica", 11)
    c.drawCentredString(width / 2, y, "[ĐƠN VỊ QUẢN LÝ HỌC PHẦN]")
    y -= 25
    c.setFont("Helvetica-Bold", 12)
    c.drawCentredString(width / 2, y, "ĐỀ THI KẾT THÚC HỌC PHẦN TRẮC NGHIỆM KHÁCH QUAN")
    y -= 25

    c.setFont("Helvetica", 10)
    c.drawString(50, y, "Học kỳ: …… Năm học: ……… - ………")
    y -= 15
    c.drawString(50, y, "Tên học phần: ...................................................  Mã học phần: ...................  Số TC: ........")
    y -= 15
    c.drawString(50, y, "Ngành: .................................................  Hệ đào tạo: ............  Hình thức đào tạo: ...............")
    y -= 30
    c.drawCentredString(width / 2, y, "ĐỀ THI GỐC")
    y -= 30

    flow.append(Paragraph("ĐỀ THI GỐC", styles['RobotoTitle']))
    flow.append(Spacer(1, 12))

    idx = 1

    if exam_type == 'mcq':
        for q in data:
            flow.append(Paragraph(f"Câu {idx}: {q['question_text']}", styles['RobotoNormal']))
            flow.append(Paragraph(f"A. {q.get('answer_a', '')}", styles['RobotoNormal']))
            flow.append(Paragraph(f"B. {q.get('answer_b', '')}", styles['RobotoNormal']))
            flow.append(Paragraph(f"C. {q.get('answer_c', '')}", styles['RobotoNormal']))
            flow.append(Paragraph(f"D. {q.get('answer_d', '')}", styles['RobotoNormal']))
            flow.append(Spacer(1, 12))
            idx += 1
    else:
        for q in data.get('MCQ', []):
            flow.append(Paragraph(f"Câu {idx}: {q['question_text']}", styles['RobotoNormal']))
            flow.append(Paragraph(f"A. {q['answer_a']}", styles['RobotoNormal']))
            flow.append(Paragraph(f"B. {q['answer_b']}", styles['RobotoNormal']))
            flow.append(Paragraph(f"C. {q['answer_c']}", styles['RobotoNormal']))
            flow.append(Paragraph(f"D. {q['answer_d']}", styles['RobotoNormal']))
            flow.append(Spacer(1, 12))
            idx += 1

        for q in data.get('tf', []):
            flow.append(Paragraph(f"Câu {idx} (Đúng/Sai): {q['question_text']}", styles['RobotoNormal']))
            if 'image_url' in q and q['image_url']:
                try:
                    flow.append(Image(q['image_url'], width=200, height=150))
                except Exception as e:
                    flow.append(Paragraph(f"[Lỗi ảnh: {str(e)}]", styles['RobotoItalic']))
            for i, st in enumerate(q['sub_items'], start=1):
                flow.append(Paragraph(f"({i}) {st['statement_text']}", styles['RobotoNormal']))
            flow.append(Spacer(1, 12))
            idx += 1

        for q in data.get('SA', []):
            flow.append(Paragraph(f"Câu {idx} (Tự luận): {q['question_text']}", styles['RobotoNormal']))
            flow.append(Spacer(1, 12))
            idx += 1

    flow.append(Paragraph("BẢNG ĐÁP ÁN", styles['RobotoHeading2']))
    flow.append(Spacer(1, 12))

    table_data = [["Câu", "Đáp án đúng"]]
    idx = 1

    if exam_type == 'mcq':
        for q in data:
            table_data.append([str(idx), q.get('correct_answer', '')])
            idx += 1
    else:
        for q in data.get('MCQ', []):
            table_data.append([str(idx), q.get('correct_answer', '')])
            idx += 1
        for q in data.get('tf', []):
            corrects = [s['statement_text'] for s in q['sub_items'] if s['is_true']]
            table_data.append([str(idx), ", ".join(corrects)])
            idx += 1
        for q in data.get('SA', []):
            table_data.append([str(idx), q.get('correct_answer', '')])
            idx += 1

    table = Table(table_data, colWidths=[50, 400])
    table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Roboto'),
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
    ]))
    flow.append(table)
    flow.append(Spacer(1, 12))
    flow.append(Paragraph("HẾT", styles['RobotoTitle']))

    doc.build(flow)
    buf.seek(0)
    return buf







## KẾT THÚC THÊM VÀO




