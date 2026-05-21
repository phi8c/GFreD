import pdfplumber
from docx import Document
import re
from werkzeug.utils import secure_filename
import PyPDF2
import os
import fitz  # PyMuPDF
from PIL import Image
import io
import uuid
from io import BytesIO

def extract_text(file_path):
    text = ""
    if file_path.endswith(".pdf"):
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
    
    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    
    else:
        return "Định dạng file không được hỗ trợ!"
    
    return text
import re


def is_toc_line(line):
    # Nhận diện dòng mục lục:
    # - chứa nhiều dấu chấm liên tiếp
    # - kết thúc bằng số
    return (
        re.search(r'\.{5,}', line) and
        re.search(r'\d+\s*$', line)
    )

def extract_text_advance(file_path):
    text = ""
    if file_path.endswith(".pdf"):
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if not page_text:
                    continue
                
                # Bỏ nguyên trang nếu có tiêu đề "mục lục"
                if re.search(r'\bmục lục\b', page_text, flags=re.IGNORECASE):
                    continue

                # Lọc từng dòng
                filtered_lines = []
                for line in page_text.splitlines():
                    if is_toc_line(line.strip().lower()):
                        continue
                    filtered_lines.append(line)
                
                text += "\n".join(filtered_lines) + "\n"

    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        for para in doc.paragraphs:
            line = para.text.strip()
            if re.search(r'\bmục lục\b', line, flags=re.IGNORECASE):
                continue
            if is_toc_line(line.lower()):
                continue
            text += line + "\n"

    else:
        return "Định dạng file không được hỗ trợ!"

    return text
def split_text_by_chapters_scd(text):
    # Tách theo chương: "Chương 1", "chương II:", "CHƯƠNG IV."...
    # Phải nằm ở đầu dòng (sát trái), có thể kèm dấu : hoặc . và chữ hoa/thường
    #pattern = r'(?i)(?=^chương\s+(?:\d+|[ivxlcdm]+)[\.:]?[^\n]*)'
    pattern = r'(?i)(?=^ch[ưuƢ][ơo]ng\s+(?:\d+|[ivxlcdm]+)[\.:]?.*$)'

    # Tách văn bản
    raw_chapters = re.split(pattern, text, flags=re.MULTILINE)

    chapters = {}
    for part in raw_chapters:
        part = part.strip()
        if not part:
            continue

        # Tìm dòng đầu tiên có từ "chương"
        lines = part.splitlines()
        #header_line = next((l for l in lines if re.match(r'(?i)^chương\s+(?:\d+|[ivxlcdm]+)', l.strip())), None)
        header_line = next((l for l in lines if re.match(r'(?i)^ch[ưuƢ][ơo]ng\s+(?:\d+|[ivxlcdm]+)', l.strip())), None)
        if not header_line:
            continue  # không hợp lệ

        # Tìm số chương (Ả Rập hoặc La Mã)
        #match = re.search(r'(?i)^chương\s+(?P<num>\d+|[ivxlcdm]+)', header_line.strip())
        match = re.search(r'(?i)^ch[ưuƢ][ơo]ng\s+(?P<num>\d+|[ivxlcdm]+)', header_line.strip())
        if not match:
            continue
        
        num_raw = match.group("num")
        try:
            # Ưu tiên chuyển La Mã → số nếu cần
            chapter_number = roman_to_int(num_raw) if is_roman(num_raw) else int(num_raw)
        except:
            continue
        
        # Nội dung là toàn bộ phần sau tiêu đề
        content_lines = lines[1:]  # bỏ dòng đầu là "chương ..."
        chapter_content = "\n".join(content_lines).strip()
        chapters[chapter_number] = chapter_content

    return chapters  # Dict[int, str]

# ====== Các hàm hỗ trợ ======

def is_roman(s):
    return re.fullmatch(r'[ivxlcdm]+', s.lower()) is not None

def roman_to_int(roman):
    roman = roman.upper()
    roman_numerals = {
        'I': 1, 'V': 5, 'X': 10,
        'L': 50, 'C': 100,
        'D': 500, 'M': 1000
    }
    total = 0
    prev = 0
    for char in reversed(roman):
        val = roman_numerals[char]
        if val < prev:
            total -= val
        else:
            total += val
            prev = val
    return total

def split_text_by_chapters(text):
    # Tách chương bằng regex: "Chương 1", "Chương 2", ...
    parts = re.split(r'(CHƯƠNG\s+\d+)', text)
    chapters = {}
    
    # Kết quả sau split là: ['', 'Chương 1', 'nội dung', 'Chương 2', 'nội dung', ...]
    for i in range(1, len(parts), 2):
        chapter_title = parts[i].strip()
        chapter_number = int(re.findall(r'\d+', chapter_title)[0])
        chapter_content = parts[i+1].strip()
        chapters[chapter_number] = chapter_content
    
    return chapters  # {1: "nội dung chương 1", 2: "nội dung chương 2", ...}

# hàm này để trích xuất số chương từ file trong giao diện chính

# ĐÂY LÀ HÀM TRÍCH XUẤT SỐ CHƯƠNG MỚI THÊM VÀO


def split_text_by_chapters_advance(text):
    import re
    
    # Chuẩn hóa text
    text = re.sub(r'[\u200b\u200c\u200d\ufeff]', '', text)  # Loại bỏ zero-width characters
    text = re.sub(r'\s+', ' ', text)  # Chuẩn hóa khoảng trắng
    text = text.strip()
    
    # Loại bỏ mục lục
    toc_patterns = [
        r'(?i)(?:^|\n)\s*(?:mục\s*lục|danh\s*mục|bảng\s*mục\s*lục|table\s*of\s*contents)\s*.*?(?=(?:ch[uư][ơo]ng\s+\d+|$))',
        r'(?i)(?:^|\n)\s*(?:contents?|index)\s*\n.*?(?=(?:ch[uư][ơo]ng\s+\d+|$))'
    ]
    
    for pattern in toc_patterns:
        text = re.sub(pattern, '', text, flags=re.DOTALL)
    
    # Loại bỏ lời nói đầu, lời mở đầu
    intro_patterns = [
        r'(?i)(?:^|\n)\s*(?:lời\s*(?:mở\s*đầu|nói\s*đầu|giới\s*thiệu|tựa)|preface|foreword|introduction)\s*.*?(?=(?:ch[uư][ơo]ng\s+\d+|$))',
        r'(?i)(?:^|\n)\s*(?:dẫn\s*nhập|khái\s*quát|tổng\s*quan|mở\s*đầu)\s*.*?(?=(?:ch[uư][ơo]ng\s+\d+|$))'
    ]
    
    for pattern in intro_patterns:
        text = re.sub(pattern, '', text, flags=re.DOTALL)
    
    # Tách chương bằng regex: "CHƯƠNG 1", "Chương 2", ...
    parts = re.split(r'(CH[UƯ][ƠO]NG\s+\d+)', text, flags=re.IGNORECASE)
    chapters = {}
    
    # Kết quả sau split là: ['', 'Chương 1', 'nội dung', 'Chương 2', 'nội dung', ...]
    for i in range(1, len(parts), 2):
        if i + 1 < len(parts):
            chapter_title = parts[i].strip()
            chapter_number = int(re.findall(r'\d+', chapter_title)[0])
            chapter_content = parts[i+1].strip()
            
            # Chỉ thêm nếu có nội dung
            if chapter_content:
                chapters[chapter_number] = chapter_content
    
    return chapters  # {1: "nội dung chương 1", 2: "nội dung chương 2", ...}




# KẾT THÚC HÀM TRÍCH XUẤT SỐ CHƯƠNG MỚI THÊM VÀO

def extract_chapters_from_file(file_path):
    # Đọc nội dung file (hỗ trợ .docx và .pdf)
    text = ""
    if file_path.endswith('.docx'):
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
    elif file_path.endswith('.pdf'):
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = "\n".join([page.extract_text() or "" for page in reader.pages])
    else:
        raise ValueError("Unsupported file type")

    # Regex tìm tất cả các chương
    chapter_pattern = re.compile(r'(Chương|CHƯƠNG|chương)\s*\d+[:：. ]?.*?(?=\n|$)')
    matches = chapter_pattern.findall(text)

    # Lấy lại toàn bộ đoạn khớp (không chỉ từ "chương")
    chapter_lines = re.findall(r'(?:Chương|CHƯƠNG|chương)\s*\d+[:：. ]?.*?(?=\n|$)', text)

    # Loại bỏ trùng, strip lại kết quả
    cleaned_chapters = list(dict.fromkeys([line.strip() for line in chapter_lines]))

    return cleaned_chapters

# kết thúc trích xuất số chương từ file trong giao diện chính
#### HÀM TRÍCH XUẤT CHƯƠNG MỚI THÊM VÀO

def merge_duplicate_chapters(chapters):
    seen_numbers = {}
    result = []
    for ch in chapters:
        match = re.search(r'chương\s+(\d+)', ch, re.IGNORECASE)
        if match:
            num = int(match.group(1))
            # Nếu chưa có chương này thì thêm
            if num not in seen_numbers:
                seen_numbers[num] = ch
                result.append(ch)
    return result

def extract_chapters_from_file_advace(file_path):
    """
    Đọc file (PDF, DOCX) và trích xuất tiêu đề các chương.
    Giữ nguyên đầu vào/đầu ra như bản gốc:
        - input: file_path (str)
        - output: list[str] (các tiêu đề chương, không trùng, giữ thứ tự)
    """

    # 1. Lấy toàn bộ text đã lọc mục lục
    text = extract_text_advance(file_path)

    # 2. Regex nhận diện tiêu đề chương (hỗ trợ cả số La Mã và Ả Rập, case-insensitive)
    chapter_pattern = re.compile(
        r'(?im)^\s*(?:Chương|CHƯƠNG|chương)\s+(?:\d+|[ivxlcdm]+)[\.:：]?[^\n]*'
    )

    # 3. Tìm tất cả dòng tiêu đề chương
    chapter_lines = chapter_pattern.findall(text)

    # 4. Loại trùng, strip khoảng trắng
    cleaned_chapters = list(dict.fromkeys(line.strip() for line in chapter_lines))

    return cleaned_chapters




##### KẾT THÚC HÀM TRÍCH XUẤT CHƯƠNG MÓI THÊM VÀO

# đây là bắt đầu đoạn mã trích xuất hình ảnh từ file

def extract_chapters_from_docx(file_path, output_image_dir):
    doc = Document(file_path)
    chapters = []
    current = {"title": None, "text": "", "images": []}
    image_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if re.match(r"(?i)^chương\s+\d+", text):  # Phát hiện tiêu đề chương
            if current["title"]:
                chapters.append(current)
                current = {"title": text, "text": "", "images": []}
            else:
                current["title"] = text
        else:
            current["text"] += text + "\n"

    # Trích xuất ảnh từ docx
    rels = doc.part._rels
    for rel in rels:
        rel = rels[rel]
        if "image" in rel.target_ref:
            image_count += 1
            img_name = f"img_{image_count:03}.png"
            image_data = rel.target_part.blob
            img_path = os.path.join(output_image_dir, img_name)
            with open(img_path, "wb") as f:
                f.write(image_data)
            current["images"].append(img_name)

    chapters.append(current)
    return chapters

def extract_chapters_from_pdf(file_path, output_image_dir):
    doc = fitz.open(file_path)

    chapters = []
    current_chapter = None
    current_text = ""
    current_images = []

    os.makedirs(output_image_dir, exist_ok=True)

    chapter_pattern = re.compile(r'(Chương|CHƯƠNG|chương)\s*\d+[:：. ]?.*?(?=\n|$)')

    for page_num, page in enumerate(doc):
        text = page.get_text()
        
        # Bỏ qua trang mục lục (nếu có tiêu đề "Mục lục", không case-sensitive)
        if re.search(r'\bMục\s+lục\b', text, re.IGNORECASE) and page_num < 3:
            print(f"⚠️ Bỏ qua trang {page_num + 1} vì là mục lục.")
            continue
        
        images = page.get_images(full=True)

        # Tìm tiêu đề chương
        #chapter_match = chapter_pattern.search(text)
        chapter_match = list(chapter_pattern.finditer(text))
        if chapter_match:
            # Lưu chương trước nếu đang có
            if current_chapter:
                chapters.append({
                    "title": current_chapter,
                    "text": current_text.strip(),
                    "images": current_images
                })

            # Bắt đầu chương mới
            current_chapter = chapter_match.group(1).strip()
            current_text = ""
            current_images = []

        current_text += text + "\n"

        for i, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            img_bytes = base_image["image"]
            ext = base_image["ext"]
            img_name = f"{uuid.uuid4().hex[:8]}.{ext}"
            img_path = os.path.join(output_image_dir, img_name)

            with open(img_path, "wb") as img_file:
                img_file.write(img_bytes)

            # Caption giả định: lấy vài dòng quanh đoạn ảnh xuất hiện
            # (có thể nâng cấp sau)
            caption = "Ảnh trích từ nội dung chương"

            current_images.append({
                "path": img_path,
                "caption": caption
            })

    # Đừng quên chương cuối
    if current_chapter:
        chapters.append({
            "title": current_chapter,
            "text": current_text.strip(),
            "images": current_images
        })

    return chapters
#chapters = extract_chapters_from_pdf("path/to/file.pdf", "output/images")


# đây là kết thúc trích xuất đoạn mã
import fitz  # PyMuPDF
import os
import re
import uuid


def extract_chapters_from_pdffs(file_path, output_image_dir):
    doc = fitz.open(file_path)

    chapters = []
    current_chapter_title = None
    current_text = ""
    current_images = []

    os.makedirs(output_image_dir, exist_ok=True)

    # Regex tìm tiêu đề chương (full dòng)
    chapter_title_pattern = re.compile(r'^\s*(Chương|CHƯƠNG|chương)\s*\d+[:：. ]?.*$', re.MULTILINE)
    
    #toc_line_pattern = re.compile(r'(Chương|CHƯƠNG|chương)\s*\d+.*?\.{3,}.*?\d{1,3}')
    toc_start = -1
    
    for i in range(min(5, len(doc))):  # chỉ kiểm tra 5 trang đầu
     text = doc[i].get_text()
     if re.search(r'\bMục\s+lục\b', text, re.IGNORECASE):
        toc_start = i
        break

    pages_to_skip = set()
    if toc_start != -1:
    # Loại trang mục lục và 2 trang tiếp theo nếu có
     pages_to_skip = set(range(toc_start, min(toc_start + 3, len(doc))))
    print(f"⚠️ Bỏ qua các trang mục lục: {sorted(pages_to_skip)}")
    
    
    for page_num, page in enumerate(doc):
        if page_num in pages_to_skip:
           continue
        text = page.get_text()
        

        # Kiểm tra có tiêu đề chương trong trang này không
        matches = list(chapter_title_pattern.finditer(text))
        if matches:
            for idx, match in enumerate(matches):
                full_chapter_title = match.group().strip()

                # Lưu chương cũ nếu đang có
                if current_chapter_title:
                    chapters.append({
                        "title": current_chapter_title,
                        "text": current_text.strip(),
                        "images": current_images
                    })

                # Reset cho chương mới
                current_chapter_title = full_chapter_title
                current_text = text[match.end():].strip() if idx == 0 else ""  # phần sau tiêu đề
                current_images = []
        else:
            current_text += "\n" + text

        # Xử lý ảnh trong trang
        for img in page.get_images(full=True):
            xref = img[0]
            base_image = doc.extract_image(xref)
            img_bytes = base_image["image"]
            ext = base_image["ext"]
            img_name = f"{uuid.uuid4().hex[:8]}.{ext}"
            img_path = os.path.join(output_image_dir, img_name)

            with open(img_path, "wb") as img_file:
                img_file.write(img_bytes)

            current_images.append({
                "path": img_path,
                "caption": "Ảnh trích từ chương"
            })

    # Lưu chương cuối cùng
    if current_chapter_title and current_text.strip():
        chapters.append({
            "title": current_chapter_title,
            "text": current_text.strip(),
            "images": current_images
        })

    print(f"✅ Tổng số chương phát hiện: {len(chapters)}")
    return chapters

#  ĐÂY LÀ CODE TRÍCH XUẤT VĂN BẢN, HÌNH ẢNH NÂNG CAO


def extract_chapters_from_pdf_advance(file_path, output_image_dir):
    """
    Trích xuất chương đơn giản - mỗi chương là một trang hoặc nhóm trang
    """
    doc = fitz.open(file_path)
    chapters = []
    
    os.makedirs(output_image_dir, exist_ok=True)
    
    # Pattern tìm tiêu đề chương
    chapter_pattern = re.compile(r'^(Chương|CHƯƠNG|chương)\s+\d+\s*[:.：-]?\s*.+$', re.MULTILINE)
    
    # Bỏ qua trang mục lục và lời nói đầu
    toc_pages = find_toc_pages(doc)
    preface_pages = find_preface_pages(doc)
    skip_pages = toc_pages.union(preface_pages)
    print(f"🔍 Bỏ qua trang mục lục: {toc_pages}")
    print(f"🔍 Bỏ qua trang lời nói đầu: {preface_pages}")
    
    current_chapter = None
    
    # Duyệt từng trang
    for page_num in range(len(doc)):
        if page_num in skip_pages:
            continue
            
        page = doc[page_num]
        page_text = clean_text(page.get_text())
        
        # Tìm tiêu đề chương trong trang này
        chapter_matches = list(chapter_pattern.finditer(page_text))
        
        if chapter_matches:
            # Lưu chương trước đó (nếu có)
            if current_chapter and len(current_chapter['text'].strip()) > 50:
                chapters.append(current_chapter)
            
            # Bắt đầu chương mới
            chapter_title = chapter_matches[0].group().strip()
            print(f"📖 Tìm thấy chương: {chapter_title}")
            
            current_chapter = {
                "title": chapter_title,
                "text": page_text,
                "images": extract_images_from_page(page, output_image_dir)
            }
        else:
            # Nếu đã có chương hiện tại, thêm text vào
            if current_chapter:
                current_chapter['text'] += "\n" + page_text
                # Chỉ thêm ảnh nếu chưa đủ 4 ảnh
                if len(current_chapter['images']) < 4:
                    page_images = extract_images_from_page(page, output_image_dir)
                    for img in page_images:
                        if len(current_chapter['images']) >= 4:
                            break
                        current_chapter['images'].append(img)
    
    # Lưu chương cuối cùng
    if current_chapter and len(current_chapter['text'].strip()) > 50:
        chapters.append(current_chapter)
    
    doc.close()
    print(f"🎉 Hoàn thành! Trích xuất {len(chapters)} chương")
    return chapters

def find_toc_pages(doc):
    """Tìm trang mục lục để bỏ qua"""
    toc_pages = set()
    
    for i in range(min(5, len(doc))):
        text = doc[i].get_text()
        if re.search(r'\b(Mục\s+lục|MỤC\s+LỤC|Table\s+of\s+Contents)\b', text, re.IGNORECASE):
            toc_pages.update(range(i, min(i + 3, len(doc))))
            break
    
    return toc_pages

def find_preface_pages(doc):
    """Tìm trang lời nói đầu để bỏ qua"""
    preface_pages = set()
    
    for i in range(min(10, len(doc))):  # Kiểm tra 10 trang đầu
        text = doc[i].get_text()
        if re.search(r'\b(LỜI\s+NÓI\s+ĐẦU|lời\s+nói\s+đầu|Lời\s+nói\s+đầu|PREFACE|Preface)\b', text, re.IGNORECASE):
            preface_pages.add(i)
            print(f"🚫 Tìm thấy trang lời nói đầu: {i+1}")
    
        return preface_pages
    """Tìm trang mục lục để bỏ qua"""
    toc_pages = set()
    
    for i in range(min(5, len(doc))):
        text = doc[i].get_text()
        if re.search(r'\b(Mục\s+lục|MỤC\s+LỤC|Table\s+of\s+Contents)\b', text, re.IGNORECASE):
            toc_pages.update(range(i, min(i + 3, len(doc))))
            break
    
    return toc_pages

def clean_text(text):
    """Làm sạch text"""
    text = re.sub(r'\x0c', '', text)  # Form feed
    text = re.sub(r'\n{3,}', '\n\n', text)  # Nhiều xuống dòng
    return text.strip()

def extract_images_from_page(page, output_dir):
    """Trích xuất ảnh từ một trang - tối đa 4 ảnh mỗi chương"""
    page_images = []
    
    try:
        for img in page.get_images(full=True):
            xref = img[0]
            doc = page.parent
            base_image = doc.extract_image(xref)
            
            img_bytes = base_image["image"]
            ext = base_image.get("ext", "png")
            img_name = f"img_{uuid.uuid4().hex[:8]}.{ext}"
            img_path = os.path.join(output_dir, img_name)
            
            # Kiểm tra kích thước ảnh (bỏ qua ảnh quá nhỏ)
            if len(img_bytes) > 1000:  # Ít nhất 1KB
                with open(img_path, "wb") as img_file:
                    img_file.write(img_bytes)
                
                page_images.append({
                    "path": img_path,
                    "caption": "Ảnh từ chương"
                })
    except Exception as e:
        print(f"⚠️ Lỗi trích xuất ảnh: {e}")
    
    return page_images

# KẾT THÚC ĐOẠN MÃ TRÍCH XUẤT VĂN BẢN, HÌNH ẢNH NÂNG CAO


###### HÀM TRÍCH XUẤT VĂN BẢN NÂNG CAO DOCX
def extract_chapters_from_docx_advance(file_path, output_image_dir):
    """
    Trích xuất chương từ tài liệu DOCX.
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"❌ Lỗi khi mở tài liệu DOCX: {e}")
        return []

    chapters = []
    os.makedirs(output_image_dir, exist_ok=True)

    chapter_pattern = re.compile(r'^(Chương|CHƯƠNG|chương)\s+\d+\s*[:.：-]?\s*.+$', re.MULTILINE)

    current_chapter = None
    all_paragraphs = doc.paragraphs
    all_pictures = []

    for rel in doc.part.rels.values():
        # Sửa lỗi: Bỏ qua mối quan hệ external
        if rel.is_external:
            continue
        
        # Tiếp tục kiểm tra nếu đó là hình ảnh nội bộ
        if rel.target_part and rel.target_part.content_type.startswith("image/"):
            try:
                image_bytes = rel.target_part.blob
                if len(image_bytes) > 1000:
                    img_ext = "png"
                    try:
                        with Image.open(BytesIO(image_bytes)) as img_check:
                            img_ext = img_check.format.lower()
                    except:
                        pass # Giữ nguyên png nếu không xác định được định dạng
                    img_name = f"img_{uuid.uuid4().hex[:8]}.{img_ext}"
                    img_path = os.path.join(output_image_dir, img_name)
                    with open(img_path, "wb") as img_file:
                        img_file.write(image_bytes)
                    all_pictures.append({"path": img_path, "caption": "Ảnh từ chương"})
            except Exception as e:
                print(f"⚠️ Lỗi trích xuất ảnh DOCX nhúng: {e}")

    picture_index = 0

    for para in all_paragraphs:
        paragraph_text = clean_text(para.text)
        chapter_matches = list(chapter_pattern.finditer(paragraph_text))

        if chapter_matches:
            if current_chapter and len(current_chapter['text'].strip()) > 50:
                chapters.append(current_chapter)
            chapter_title = chapter_matches[0].group().strip()
            current_chapter = {
                "title": chapter_title,
                "text": paragraph_text,
                "images": []
            }
        else:
            if current_chapter:
                current_chapter['text'] += "\n" + paragraph_text

        if current_chapter and len(current_chapter['images']) < 4 and picture_index < len(all_pictures):
            current_chapter['images'].append(all_pictures[picture_index])
            picture_index += 1

    if current_chapter and len(current_chapter['text'].strip()) > 50:
        chapters.append(current_chapter)

    return chapters

##### KẾT THÚC


###### TRÍCH XUẤT SỐ LƯỢNG CHƯƠNG TRONG FILE

def extract_chapters_from_file_advance(file_path):
    def clean_text(text):
        text = re.sub(r'\x0c', '', text)  # Form feed
        text = re.sub(r'\n{3,}', '\n\n', text)  # Nhiều xuống dòng
        return text.strip()

    def find_skip_pages_pdf(reader):
        toc_pages = set()
        preface_pages = set()

        for i in range(min(10, len(reader.pages))):
            text = reader.pages[i].extract_text() or ""
            if re.search(r'\b(Mục\s+lục|MỤC\s+LỤC|Table\s+of\s+Contents)\b', text, re.IGNORECASE):
                toc_pages.update(range(i, min(i + 3, len(reader.pages))))
            if re.search(r'\b(LỜI\s+NÓI\s+ĐẦU|lời\s+nói\s+đầu|PREFACE|Preface)\b', text, re.IGNORECASE):
                preface_pages.add(i)
        return toc_pages.union(preface_pages)

    text = ""
    if file_path.endswith('.docx'):
        doc = Document(file_path)
        paras = [para.text for para in doc.paragraphs if para.text.strip()]
        joined_text = "\n".join(paras)
        text = clean_text(joined_text)

    elif file_path.endswith('.pdf'):
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            skip_pages = find_skip_pages_pdf(reader)
            texts = []
            for i, page in enumerate(reader.pages):
                if i in skip_pages:
                    continue
                page_text = page.extract_text() or ""
                texts.append(page_text)
            text = clean_text("\n".join(texts))
    else:
        raise ValueError("Unsupported file type")

    # Pattern tìm tiêu đề chương, hỗ trợ số hoặc số La Mã
    chapter_pattern = re.compile(
        r'(?m)^(Chương|CHƯƠNG|chương)\s+(?:\d+|[IVXLCDM]+)\s*[:.：-]?\s*.*$'
    )

    matches = list(chapter_pattern.finditer(text))

    chapter_lines = []
    for m in matches:
        line = m.group().strip()
        if line not in chapter_lines:
            chapter_lines.append(line)

    return chapter_lines

##### KẾT THÚC TRÍCH XUẤT SỐ LƯỢNG CHƯƠNG TRONG FILE


