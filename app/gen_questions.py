from flask import Flask, logging, request, render_template, jsonify,  Blueprint, current_app, flash
from extract_text import extract_text, split_text_by_chapters, extract_chapters_from_file, split_text_by_chapters_advance, extract_text_advance, split_text_by_chapters_scd, extract_chapters_from_file_advace
from generates_questions import generate_questions, generate_questions_from_jobs, generate_questions_advance
import os
import mysql.connector
from flask import send_file
from docx import Document
import json
from flask import Flask
from datetime import datetime, timedelta
from flask_socketio import SocketIO
from io import BytesIO
from flask import Blueprint, render_template, request, redirect, url_for
from werkzeug.utils import secure_filename
import re
import string
import random 
from werkzeug.security import check_password_hash, generate_password_hash

# khai báo những gì cần thiết cho đăng nhập
from flask_login import LoginManager, current_user, login_required
from models.user_model import User
from controllers.db_users import get_user_by_id
#from utils.export_exam_utils import render_exam_to_docx, render_exam_to_pdf

from routes.auth_routes import auth_bp  # <-- Blueprint
from sentence_transformers import SentenceTransformer, util



from dotenv import load_dotenv

load_dotenv() 

question_bp = Blueprint('question', __name__)

model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
models = SentenceTransformer('paraphrase-multilingual-mpnet-base-v2')
now = datetime.now()




if not os.path.exists('uploads'):
    os.makedirs('uploads')
difficulty_map = {
    "de": 0,
    "trung_binh": 1,
    "kho": 2
}
@login_required
@question_bp.route('/extract_chapters', methods=['POST'])
def extract_chapters():
    file = request.files.get('file')
    if not file:
        return jsonify({'error': 'No file provided'}), 400

    filename = secure_filename(file.filename)
    file_pathh = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
    file.save(file_pathh)

    try:
        print("📥 Nhận file và bắt đầu tách chương")
        chapters = extract_chapters_from_file_advace(file_pathh)
        print("📚 Danh sách chương trích xuất:", chapters)
        return jsonify({'chapters': chapters})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
def is_duplicate_question(new_question: str, exclude_id=None, threshold=0.85):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # Lấy tất cả câu hỏi hiện tại trong DB
    sql = "SELECT id, question_text FROM question_bank"
    if exclude_id:  # khi update thì bỏ qua chính nó
        sql += " WHERE id != %s"
        cursor.execute(sql, (exclude_id,))
    else:
        cursor.execute(sql)

    rows = cursor.fetchall()
    cursor.close()
    conn.close()

    if not rows:
        return False, None

    # Encode câu hỏi mới
    new_embedding = models.encode(new_question, convert_to_tensor=True)

    # Kiểm tra trùng với từng câu trong DB
    for row in rows:
        existing_embedding = models.encode(row["question_text"], convert_to_tensor=True)
        similarity = util.cos_sim(new_embedding, existing_embedding).item()
        if similarity >= threshold:
            return True, row["id"]

    return False, None
    
    
@question_bp.route('/check-exam-code/<int:exam_set_id>')
def check_exam_code(exam_set_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("""
        SELECT COUNT(*) AS cnt
        FROM exam_codes
        WHERE exam_set_id = %s
    """, (exam_set_id,))
    result = cursor.fetchone()

    cursor.close()
    conn.close()
    
    #print("in ra reusult", result)

    return jsonify({"has_code": result['cnt'] > 0})

    
    
    
# ROUTE NÀY MỚI THÊM VÀO ĐỂ TẠO NGÂN HÀNG CÂU HỎI
@question_bp.route("/check-exam-set-in-bank/<int:exam_set_id>", methods=["GET"])
def check_exam_set_in_bank(exam_set_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # Kiểm tra có câu hỏi trong ngân hàng từ exam_set_id
    cursor.execute("""
        SELECT 1 FROM question_bank 
        WHERE exam_set_id = %s LIMIT 1
    """, (exam_set_id,))
    in_bank = cursor.fetchone() is not None

    # Hoặc kiểm tra nếu bộ đề được tạo từ ngân hàng (sources = 'from_bank')
    if not in_bank:
        cursor.execute("""
            SELECT sources FROM exam_sets WHERE id = %s
        """, (exam_set_id,))
        row = cursor.fetchone()
        if row and row.get("sources") == "from_bank":
            in_bank = True

    cursor.close()
    conn.close()
    return jsonify({"exists": in_bank})


@question_bp.route("/generate-from-bank", methods=["GET", "POST"])
@login_required
def generate_from_bank():
    user_id = current_user.id
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    if request.method == "POST":
        num_questions = int(request.form.get("num_questions"))
        subject_name = request.form.get("subject_name")

        # Lấy tất cả câu hỏi trong ngân hàng theo môn học
        cursor.execute("""
            SELECT * FROM question_bank
            WHERE user_id = %s AND subject_name = %s
        """, (user_id, subject_name))
        available_questions = cursor.fetchall()

        if len(available_questions) < num_questions:
            cursor.close()
            conn.close()
            return f"❌ Không đủ câu hỏi trong ngân hàng (hiện có {len(available_questions)}).", 400

        # Tạo bộ đề mới
        created_at = datetime.now()
        set_name = f"Bộ đề từ ngân hàng - {subject_name}"
        sources = "from_bank"

        cursor.execute("""
            INSERT INTO exam_sets (set_name, user_id, subject_name, sources, created_at)
            VALUES (%s, %s, %s, %s, %s)
        """, (set_name, user_id, subject_name, sources, created_at))
        exam_set_id = cursor.lastrowid

        # Chọn ngẫu nhiên câu hỏi
        selected_questions = random.sample(available_questions, num_questions)

        for q in selected_questions:
            cursor.execute("""
                INSERT INTO questions (
                    exam_set_id,
                    question_text,
                    answer_a,
                    answer_b,
                    answer_c,
                    answer_d,
                    correct_answer,
                    difficulty,
                    chapter
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (
                exam_set_id,
                q["question_text"],
                q["answer_a"],
                q["answer_b"],
                q["answer_c"],
                q["answer_d"],
                q["correct_answer"],
                q.get("level", 0),  # mapped từ level
                None  # chapter để trống
            ))

        conn.commit()
        cursor.close()
        conn.close()

        return redirect(url_for("question.edit_questions", exam_set_id=exam_set_id))

    # GET method: render giao diện chọn môn và số lượng
    cursor.execute("""
        SELECT DISTINCT subject_name FROM question_bank
        WHERE user_id = %s
    """, (user_id,))
    subjects = [row["subject_name"] for row in cursor.fetchall()]

    cursor.close()
    conn.close()

    return render_template("generate_from_bank.html", subjects=subjects)



@question_bp.route("/add-exam-set-to-bank/<int:exam_set_id>", methods=["POST"])
@login_required
def add_exam_set_to_bank(exam_set_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # 🟢 Lấy thông tin bộ đề
    cursor.execute("""
        SELECT id, subject_name, user_id FROM exam_sets WHERE id = %s
    """, (exam_set_id,))
    exam_set = cursor.fetchone()

    if not exam_set:
        cursor.close()
        conn.close()
        return jsonify({"error": "Không tìm thấy bộ đề"}), 404

    user_id = current_user.id
    subject_name = exam_set["subject_name"]

    # 🟢 Kiểm tra nếu đã có trong ngân hàng
    cursor.execute("""
        SELECT 1 FROM question_bank WHERE exam_set_id = %s LIMIT 1
    """, (exam_set_id,))
    if cursor.fetchone():
        cursor.close()
        conn.close()
        return jsonify({"message": "Bộ đề đã có trong ngân hàng"}), 200

    # 🟢 Lấy câu hỏi thuộc bộ đề này
    cursor.execute("""
        SELECT *
        FROM questions
        WHERE exam_set_id = %s
    """, (exam_set_id,))
    exam_questions = cursor.fetchall()

    if not exam_questions:
        cursor.close()
        conn.close()
        return jsonify({"message": "Không có câu hỏi để thêm vào ngân hàng"}), 200

    # 🟢 Lấy câu hỏi trong ngân hàng cùng môn để so sánh
    cursor.execute("""
        SELECT question_text FROM question_bank
        WHERE user_id = %s AND subject_name = %s
    """, (user_id, subject_name))
    existing_questions = [row["question_text"] for row in cursor.fetchall()]
    existing_embeddings = models.encode(existing_questions, convert_to_tensor=True) if existing_questions else []

    added = 0
    skipped = 0

    for q in exam_questions:
        question_text = q["question_text"]
        question_embedding = models.encode(question_text, convert_to_tensor=True)

        is_duplicate = False
        if existing_embeddings:
            cosine_scores = util.cos_sim(question_embedding, existing_embeddings)
            if float(cosine_scores.max()) > 0.85:
                is_duplicate = True

        if is_duplicate:
            skipped += 1
            continue

        level = q.get("difficulty", 0) or 0

        cursor.execute("""
            INSERT IGNORE INTO question_bank (
                question_id,
                exam_set_id,
                user_id,
                subject_name,
                level,
                question_text,
                answer_a,
                answer_b,
                answer_c,
                answer_d,
                correct_answer,
                created_at
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            q["id"],
            exam_set_id,
            user_id,
            subject_name,
            level,
            q["question_text"],
            q["answer_a"],
            q["answer_b"],
            q["answer_c"],
            q["answer_d"],
            q["correct_answer"],
            datetime.now()
        ))

        added += 1

    conn.commit()
    cursor.close()
    conn.close()

    return jsonify({
        "message": "Đã thêm vào ngân hàng",
        "added": added,
        "skipped": skipped,
        "total": len(exam_questions),
        "subject_name": subject_name
    })

### KIỂM TRA TRÙNG TRONG BỘ ĐỀ

@question_bp.route("/check-duplicate-questions/<int:exam_set_id>", methods=["GET"])
def check_duplicate_questions(exam_set_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # Lấy câu hỏi trong bộ đề
    cursor.execute("""
        SELECT id, question_text
        FROM questions
        WHERE exam_set_id = %s
    """, (exam_set_id,))
    questions = cursor.fetchall()

    cursor.close()
    conn.close()

    if len(questions) < 2:
        return jsonify({"message": "Bộ đề không đủ câu hỏi để so sánh.", "duplicates": []})

    # ✅ Dùng mô hình mạnh hơn
    #models = SentenceTransformer('paraphrase-multilingual-mpnet-base-v2')

    texts = [q["question_text"] for q in questions]
    embeddings = models.encode(texts, convert_to_tensor=True)

    # So sánh tất cả các cặp
    duplicates = []
    threshold = 0.85
    for i in range(len(questions)):
        for j in range(i + 1, len(questions)):
            score = float(util.cos_sim(embeddings[i], embeddings[j]))
            if score > threshold:
                duplicates.append({
                    "question_1": questions[i]["question_text"],
                    "question_2": questions[j]["question_text"],
                    "similarity": round(score, 4),
                    "id_1": questions[i]["id"],
                    "id_2": questions[j]["id"]
                })

    return jsonify({
        "exam_set_id": exam_set_id,
        "total": len(questions),
        "duplicates": duplicates,
        "count": len(duplicates),
        "message": f"Phát hiện {len(duplicates)} cặp câu hỏi nghi ngờ trùng."
    })



#### KẾT THÚC KIỂM TRA TRÙNG TRONG BỘ ĐỀ


    
@question_bp.route("/question-bank", methods=["GET"])
@login_required
def view_question_bank():
    user_id = current_user.id
    subject_param = request.args.get("subject")

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # 🟢 Lấy danh sách các môn học đã có trong ngân hàng
    cursor.execute("""
        SELECT DISTINCT subject_name 
        FROM question_bank
        WHERE user_id = %s
    """, (user_id,))
    rows = cursor.fetchall()
    subjects = [row['subject_name'] for row in rows]

    if not subjects:
        cursor.close()
        conn.close()
        return render_template("question_bank.html", subjects=[], selected_subject=None, questions=[])

    selected_subject = subject_param if subject_param in subjects else subjects[0]

    # 🟢 Lấy trực tiếp các câu hỏi từ bảng question_bank
    cursor.execute("""
        SELECT id, question_text, answer_a, answer_b, answer_c, answer_d, correct_answer, level
        FROM question_bank
        WHERE user_id = %s AND subject_name = %s
    """, (user_id, selected_subject))
    questions = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template(
        "question_bank.html",
        subjects=subjects,
        selected_subject=selected_subject,
        questions=questions
    )







###### KẾT THÚC ROUTE

##### route này để lấy danh sách môn học từ bảng bộ đề
@question_bp.route("/generate-subject", methods=["GET"])
def generate_subject():
    conn = get_db_connection()
    cursor = conn.cursor()

    # Lấy danh sách môn học, loại bỏ trùng lặp
    cursor.execute("""
        SELECT DISTINCT subject_name
        FROM exam_sets
        WHERE subject_name IS NOT NULL AND subject_name <> ''
    """)
    subjects = [row[0] for row in cursor.fetchall()]

    cursor.close()
    conn.close()

    # Render template và truyền danh sách môn học
    return render_template("question_bank.html", subjects=subjects)



### kết thúc route lấy danh sách môn học từ bảng bộ đề

##### route này thêm vào đê thêm câu hỏi vào ngân hàng

@question_bp.route("/api/add-question", methods=["POST"])
def add_question():
    data = request.get_json()

    # Lấy dữ liệu từ request
    subject_name = data.get("subject_name")
    level = data.get("level")
    question_text = data.get("question_text")
    answer_a = data.get("answer_a")
    answer_b = data.get("answer_b")
    answer_c = data.get("answer_c")
    answer_d = data.get("answer_d")
    correct_answer = data.get("correct_answer")

    # user_id có thể lấy từ session hoặc mặc định
    user_id = current_user.id  # mặc định 1 là admin

    # Validate dữ liệu
    if not all([subject_name, level, question_text, answer_a, answer_b, answer_c, answer_d, correct_answer]):
        return jsonify({"success": False, "message": "Thiếu dữ liệu bắt buộc."}), 400
    
    is_dup, dup_id = is_duplicate_question(question_text)
    if is_dup:
        return jsonify({
            "success": False,
            "message": f"Câu hỏi trùng với câu có id = {dup_id}"
        }), 400

    # Kết nối DB
    conn = get_db_connection()
    cursor = conn.cursor()

    try:
        cursor.execute("""
            INSERT INTO question_bank (
                question_id, exam_set_id, user_id, subject_name, level,
                answer_a, answer_b, answer_c, answer_d, correct_answer,
                question_text, created_at
            )
            VALUES (
                NULL, NULL, %s, %s, %s,
                %s, %s, %s, %s, %s,
                %s, NOW()
            )
        """, (
            user_id, subject_name, level,
            answer_a, answer_b, answer_c, answer_d, correct_answer,
            question_text
        ))

        conn.commit()
    except Exception as e:
        conn.rollback()
        return jsonify({"success": False, "message": f"Lỗi: {str(e)}"}), 500
    finally:
        cursor.close()
        conn.close()

    return jsonify({"success": True, "message": "Thêm câu hỏi thành công."})


### kết thúc route thêm câu hỏi vào ngân hàng

###### các route mối thêm vào

@question_bp.route("/question-bank/<int:id>")
def get_question(id):
    conn = get_db_connection()
    cur = conn.cursor(dictionary=True)
    cur.execute("SELECT * FROM question_bank WHERE id=%s", (id,))
    question = cur.fetchone()
    conn.close()
    print("in ra question", question)
    return jsonify(question)

@question_bp.route("/update-question-bank", methods=["POST"])
def update_question_bank():
    
    print("Form data:", request.form)
    qid = request.form["id"]
    content = request.form["content"]
    option_a = request.form["option_a"]
    option_b = request.form["option_b"]
    option_c = request.form["option_c"]
    option_d = request.form["option_d"]
    correct_answer = request.form["correct_answer"]
    difficulty = request.form["difficulty"]
    
    
    
    is_dup, dup_id = is_duplicate_question(content)
    if is_dup:
        return jsonify({
            "success": False,
            "message": f"Câu hỏi trùng với câu có id = {dup_id}"
        }), 400

    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("""
        UPDATE question_bank
        SET question_text=%s, answer_a=%s, answer_b=%s, answer_c=%s, answer_d=%s,
            correct_answer=%s, level=%s
        WHERE id=%s
    """, (content, option_a, option_b, option_c, option_d, correct_answer,difficulty, qid))
    conn.commit()
    conn.close()

    return jsonify({"success": True})

@question_bp.route("/question-bank-list")
def question_bank_list():
    conn = get_db_connection()
    cur = conn.cursor(dictionary=True)
    cur.execute("SELECT * FROM question_bank")
    rows = cur.fetchall()
    conn.close()
    return render_template("question_bank.html", questions=rows)

@question_bp.route("/delete-question/<int:question_id>", methods=["DELETE"])
def delete_question_bank(question_id):
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("DELETE FROM question_bank WHERE id = %s", (question_id,))
    conn.commit()
    cur.close()
    conn.close()
    return jsonify({"status": "success"})


#### kết thúc các route mới thêm vào
    
# route này mới thêm vào
@question_bp.route('/exam_set/<int:exam_set_id>', methods=['GET'])
@login_required
def view_exam_set(exam_set_id):
    user_id = current_user.id
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    
    
    mcq_versions = {}
    tf_versions = {}
    sa_versions = {}

    if not has_exam_codes(exam_set_id):
    # CHƯA chia mã đề → lấy đề gốc
        if is_full_thpt2025_exam_set(exam_set_id):
        # Đề gốc dạng THPT 2025
            original_data = get_original_questions_thpt2025(exam_set_id)
            #print("in ra đề gốc dạng thpt 2025", original_data)

        # Nhóm lại giống như các version theo key 'original'
            mcq_versions['original'] = {
            'exam_code_id': None,
            'questions': [
                {
                    'question_id': q['id'],
                    'question_text': q['question_text'],
                    'difficulty': q.get('difficulty'),
                    'answer_a': q.get('answer_a'),
                    'answer_b': q.get('answer_b'),
                    'answer_c': q.get('answer_c'),
                    'answer_d': q.get('answer_d'),
                    'correct_answer': q.get('correct_answer')
                } for q in original_data['MCQ']
            ]
        }  
            tf_versions['original'] = []
            for idx, q in enumerate(original_data['tf']):
                raw_items = q['sub_items']
                processed_statements = []
                for j, item in enumerate(raw_items):
                    processed_statements.append({
                    'label': string.ascii_uppercase[j],  # A, B, C...
                    'content': item['statement_text'],
                    'is_true': bool(item['is_true'])
                })

                tf_versions['original'].append({
                'id': idx,
                'question_text': q['question_text'],
                 'image_url': q.get('image_url'),  # ✅ Lấy ảnh từ dict gốc
                'statements': processed_statements,
                'position': idx + 1
                })
            # tf_versions['original'] = [
            # {
            #     'id': i,
            #     'question_text': q['question_text'],
            #     'statements': q['sub_items'],
            #     'position': idx + 1
            # }
            #for idx, (i, q) in enumerate({q['question_text']: q for q in original_data['tf']}.items())
        #]

            sa_versions['original'] = [
            {
                'id': q['id'],
                'question_text': q['question_text'],
                'position': idx + 1,
                'correct_answer': q.get('correct_answer')
            }
            for idx, q in enumerate(original_data['SA'])
        ]

            is_thpt_2025 = True

        else:
        # Đề gốc dạng MCQ
            #if is_advanced_mcq_exam_set(exam_set_id):
            mcq_data = get_mcq_original_questions_by_exam_set(exam_set_id)
            #print("in ra mcq gốc", mcq_data)
            

            mcq_versions['original'] = {
                'exam_code_id': None,
                'questions': mcq_data
            }

            is_thpt_2025 = False

    else:
    # ĐÃ chia mã đề → giữ nguyên phần mã cũ bạn đã viết
    # ---------------------- MCQ versions ----------------------
        cursor.execute(""" 
        SELECT 
            q.id AS question_id,
            q.question_text AS original_question_text,
            q.difficulty,
            eqm.exam_code_id,
            eqm.answer_a,
            eqm.answer_b,
            eqm.answer_c,
            eqm.answer_d,
            eqm.correct_answer,
            ec.code
        FROM questions q
        LEFT JOIN exam_question_map eqm ON eqm.question_id = q.id
        LEFT JOIN exam_codes ec ON ec.id = eqm.exam_code_id
        WHERE q.user_id = %s AND q.exam_set_id = %s
        ORDER BY q.id, eqm.exam_code_id
    """, (current_user.id, exam_set_id))
        rows = cursor.fetchall()

        for row in rows:
            code = row['code'] or 'original'
            if code not in mcq_versions:
                mcq_versions[code] = {
                'exam_code_id': row['exam_code_id'],
                'questions': []
            }
            mcq_versions[code]['questions'].append({
            'question_id': row['question_id'],
            'question_text': row['original_question_text'],
            'difficulty': row['difficulty'],
            'answer_a': row['answer_a'],
            'answer_b': row['answer_b'],
            'answer_c': row['answer_c'],
            'answer_d': row['answer_d'],
            'correct_answer': row['correct_answer']
        })

    # ---------------------- TF & SA versions (THPT 2025) ----------------------
        is_thpt_2025 = is_full_thpt2025_exam_set(exam_set_id)
        if is_thpt_2025:
            cursor.execute("SELECT id, code FROM exam_codes WHERE exam_set_id = %s ORDER BY code", (exam_set_id,))
            exam_codes = cursor.fetchall()

            for ec in exam_codes:
                code = ec['code']
                ec_id = ec['id']

                cursor.execute("""
                SELECT hqm.id, hqm.question_type, hqm.original_question_id, hqm.position
                FROM hs_question_map hqm
                WHERE hqm.exam_code_id = %s
                ORDER BY hqm.position
            """, (ec_id,))
                mappings = cursor.fetchall()

                for m in mappings:
                    #print(f"--> question_type = [{m['question_type']}]")

                    if m['question_type'] == 'tf':
                        if code not in tf_versions:
                            tf_versions[code] = []

                        cursor.execute("SELECT * FROM tf_questions WHERE id = %s", (m['original_question_id'],))
                        q = cursor.fetchone()
                        if q:
                            cursor.execute("""
                            SELECT label, content, is_true, position
                            FROM hs_tf_statements
                            WHERE hs_map_id = %s
                            ORDER BY position
                        """, (m['id'],))
                        statements = cursor.fetchall()
                        q['statements'] = statements
                        tf_versions[code].append({'position': m['position'], **q})
                        
                   
                    elif m['question_type'] == '':
                        #print("chạy tới điều kiện SA")
                        if code not in sa_versions:
                            #print(" chạy vào nếu không có code")
                            sa_versions[code] = []

                        cursor.execute("SELECT * FROM short_answer_questions WHERE id = %s", (m['original_question_id'],))
                        #print(" chạy vào truy vấn")
                        q = cursor.fetchone()
                        if q:
                            #print(f"[SA] Mã đề: {code}, Vị trí: {m['position']}, Câu hỏi: {q['question_text']}")
                            sa_versions[code].append({'position': m['position'], **q})
                            
                            
    #print("in ra sa question", sa_versions)
    #print("in ra dạng mcq", mcq_versions)
    print("in ra dạng đúng sai", tf_versions)
                            
    cursor.close()
    conn.close()

    return render_template(
        'exam_set_detail.html',
        exam_set_id=exam_set_id,
        mcq_versions=mcq_versions,
        tf_versions=tf_versions,
        sa_versions=sa_versions,
        is_thpt_2025=is_thpt_2025,
        
    )
    
    
    
    ## Bắt đầu đóng route lớn

    # query = """
    #     SELECT 
    #     q.id AS question_id,
    #     q.question_text AS original_question_text,
    #     q.difficulty,
    #     eqm.exam_code_id,
    #     eqm.answer_a,
    #     eqm.answer_b,
    #     eqm.answer_c,
    #     eqm.answer_d,
    #     eqm.correct_answer,
    #     ec.code
    # FROM 
    #     questions q
    # LEFT JOIN 
    #     exam_question_map eqm ON eqm.question_id = q.id
    # LEFT JOIN 
    #     exam_codes ec ON ec.id = eqm.exam_code_id
    # WHERE 
    #     q.user_id = %s AND q.exam_set_id = %s
    # ORDER BY 
    #     q.id, eqm.exam_code_id;
    # """

    # cursor.execute(query, (user_id, exam_set_id))
    # rows = cursor.fetchall()
    
    
    # exam_versions = {}

    # for row in rows:
    #     code = row['code']
    #     if code not in exam_versions:
    #         exam_versions[code] = {
    #         'exam_code_id': row['exam_code_id'],
    #         'questions': []
    #     }
    #     exam_versions[code]['questions'].append({
    #     'question_id': row['question_id'],
    #     'question_text': row['original_question_text'],
    #     'difficulty': row['difficulty'],
    #     'answer_a': row['answer_a'],
    #     'answer_b': row['answer_b'],
    #     'answer_c': row['answer_c'],
    #     'answer_d': row['answer_d'],
    #     'correct_answer': row['correct_answer']
    # })

    # # -------------------------------
    # # PHẦN 2: Kiểm tra và xử lý bộ đề THPT 2025 (nếu có)
    # # -------------------------------
    # cursor.execute("SELECT COUNT(*) as count FROM tf_questions WHERE exam_set_id = %s", (exam_set_id,))
    # is_tf = cursor.fetchone()['count'] > 0
    # cursor.execute("SELECT COUNT(*) as count FROM short_answer_questions WHERE exam_set_id = %s", (exam_set_id,))
    # is_sa = cursor.fetchone()['count'] > 0
    # cursor.execute("SELECT COUNT(*) as count FROM questions WHERE exam_set_id = %s", (exam_set_id,))
    # is_mcq_2025 = cursor.fetchone()['count'] > 0

    # is_thpt_2025 = is_tf or is_sa or is_mcq_2025
    # thpt_data = {}
    
    # tf_versions = {}
    # sa_versions = {}
    
    # if is_thpt_2025:
    #     # Lấy các mã đề của bộ đề này
    #     cursor.execute("SELECT id, code FROM exam_codes WHERE exam_set_id = %s ORDER BY code", (exam_set_id,))
    #     exam_codes = cursor.fetchall()

    #     for ec in exam_codes:
    #         code = ec['code']
    #         ec_id = ec['id']

    #         # Lấy danh sách câu hỏi từ hs_question_map
    #         cursor.execute("""
    #             SELECT hqm.id, hqm.question_type, hqm.original_question_id, hqm.position
    #             FROM hs_question_map hqm
    #             WHERE hqm.exam_code_id = %s
    #             ORDER BY hqm.position
    #         """, (ec_id,))
    #         mappings = cursor.fetchall()

    #         for m in mappings:
    #             if m['question_type'] == 'tf':
    #                 if code not in tf_versions:
    #                     tf_versions[code] = []

    #                 cursor.execute("SELECT * FROM tf_questions WHERE id = %s", (m['original_question_id'],))
    #                 q = cursor.fetchone()
    #                 if q:
    #                     cursor.execute("""
    #                         SELECT label, content, is_true, position
    #                         FROM hs_tf_statements
    #                         WHERE hs_map_id = %s
    #                         ORDER BY position
    #                     """, (m['id'],))
    #                     statements = cursor.fetchall()
    #                     q['statements'] = statements
    #                     tf_versions[code].append({'position': m['position'], **q})

    #             elif m['question_type'] == ' ':
    #                 if code not in sa_versions:
    #                     sa_versions[code] = []

    #                 cursor.execute("SELECT * FROM short_answer_questions WHERE id = %s", (m['original_question_id'],))
    #                 q = cursor.fetchone()
    #                 if q:
    #                     sa_versions[code].append({'position': m['position'], **q})

    # cursor.close()
    # conn.close()

    # return render_template(
    #     'exam_set_detail.html',
    #     exam_set_id=exam_set_id,
    #     mcq_versions=exam_versions,
    #     tf_versions=tf_versions,
    #     sa_versions=sa_versions,
    #     is_thpt_2025=is_thpt_2025
    # )
    
    ###### Kết thúc ddongsw route lớn


    

    # return render_template(
    #     'exam_set_detail.html',
    #     questions=list(questions.values()),  # MCQ truyền thống
    #     exam_set_id=exam_set_id,
    #     thpt_data=thpt_data                   # Thêm data bộ đề THPT 2025 (nếu có)
    # )
    

    # cursor.close()
    # conn.close()

    # return render_template('exam_set_detail.html', questions=list(questions.values()), exam_set_id=exam_set_id)

# kết thúc route mới thêm vào

    
@question_bp.route("/", methods=["GET", "POST"])
def index():
    questions = []
    
    user_id = current_user.id
    exam_sets = get_exam_sets_by_user(user_id)
    
    #print("in ra danh sách bộ đề của ngừi dùng", exam_sets)
    
    print("in ra id người dùng đã đăng nhập", user_id)
    
    #return "✅ Đã nhận đủ dữ liệu POST! Không gọi sinh câu hỏi lúc này."
    if request.method == "POST":
        if "file" not in request.files:
            return "Không tìm thấy file!"
        
        
        file = request.files["file"]
        if file.filename == "":
            return "Chưa chọn file!"
        
        file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        
        
        # Trích xuất văn bản từ file PDF/Word
        extracted_text = extract_text_advance(file_path)
        
        #print("Văn bản trích từ file:", extracted_text)
   
        # Gọi API OpenAI sinh câu hỏi
        mode = request.form.get("mode", "basic")
        config = request.form.get("config")
        default_name = f"Bộ đề {now.day}_{now.month}_{now.year}_{now.hour}_{now.minute}"
        set_name = request.form.get("set_name") or default_name
        print("Mode nhận được first:", mode)
        print("Config nhận được first:", config)
        print("📥 File nhận được:", file.filename)
        print("📥 Mode nhận được:", mode)
        print("📥 Config nhận được:", config)
        
        
        

        # Ngắt ở đây để kiểm tra form submit trước, không gọi GPT
        #return "✅ Đã nhận đủ dữ liệu POST! Không gọi sinh câu hỏi lúc này."

        
        # Tạo exam_set mới
        set_id = create_exam_set(set_name, user_id = user_id)

        if mode == "advanced" and config:
            config = json.loads(config)
            #jobs = []
            
            # đoạn này mới thêm vào
            
            selected_chapters = config.get('chapters_selected', [])
    
          # Truyền selected_chapters xuống hàm sinh câu hỏi
            #questions = generate_questions_from_selected_chapters(file_path, selected_chapters, total_questions)
            
            # kết thúc đoạn mới thêm vào
            chapter_texts = split_text_by_chapters_scd(extracted_text)
            # Sau khi không tách được chương
            if not chapter_texts:
                flash("❗Không thể phát hiện các chương trong tài liệu. Vui lòng kiểm tra tiêu đề chương như 'Chương 1', 'CHƯƠNG I'... mỗi chương nên bắt đầu ở dòng mới.", "warning")
                return redirect(url_for("index"))  # về lại giao diện ban đầu

            print("chapter_texts in ra:", chapter_texts)
            all_questions = []
            #all_jobs = []
            print("Mode nhận được tw:", mode)
            print("Config nhận được tw:", config)
            #return "chạy đến trước vòng lặp"

            for item in config["cau_hoi_theo_chuong"]:
                jobs = []
                chuong = item["chuong"]
                if item.get("de", 0) > 0:
                    jobs.append({"chuong": chuong, "do_kho": "de", "so_luong": item["de"]})
                if item.get("trung_binh", 0) > 0:
                    jobs.append({"chuong": chuong, "do_kho": "trung_binh", "so_luong": item["trung_binh"]})
                if item.get("kho", 0) > 0:
                    jobs.append({"chuong": chuong, "do_kho": "kho", "so_luong": item["kho"]})
                    print("⚠️ In ra job ở phần gọi Ai sinh câu hỏi!", jobs)
                if jobs and chuong in chapter_texts:
                 chapter_content = chapter_texts[chuong]
                 print("⚠️ In ra job ở phần gọi Ai sinh câu hỏi!", chapter_content)
                 print(f"📤 Gửi jobs cho chương: {chuong}", jobs)
                 
                 try:

                   questions = generate_questions_from_jobs(chapter_content, jobs)
                   all_questions.extend(questions)
                   #print("⚠️ In ra job ở phần gọi Ai sinh câu hỏi!", jobs)
                 except Exception as e:
                   print(f"❌ Lỗi khi sinh câu hỏi cho chương {chuong}: {e}")
                else:
                  print(f"⚠️ Không tìm thấy nội dung cho chương {chuong} hoặc không có jobs.")
        else:
            #questions = generate_questions(extracted_text)
            # đoạn này mói thêm vào
            #config = json.loads(config)
            
            print("➡️ Đang xử lý chế độ THƯỜNG")

            num_questions = int(request.form.get("tong_so_cau", 50))  # lấy tổng số câu hỏi
            print("Tổng số câu hỏi cầu tạo", num_questions)
            #return "Tổng số câu hỏi cần tạo"
            chapter_texts = split_text_by_chapters_scd(extracted_text)
            
            if not chapter_texts:
                flash("❗Không thể phát hiện các chương trong tài liệu. Vui lòng kiểm tra tiêu đề chương như 'Chương 1', 'CHƯƠNG I'... mỗi chương nên bắt đầu ở dòng mới.", "warning")
                return redirect(url_for("index"))  # về lại giao diện ban đầu
            chapter_count = len(chapter_texts)

            if chapter_count == 0:
              print("🚨 Không tìm thấy chương nào trong văn bản!")
              return "Không có chương hợp lệ", 400

            base = num_questions // chapter_count
            remainder = num_questions % chapter_count
            distribution = [base + 1 if i < remainder else base for i in range(chapter_count)]

            all_questions = []
            for idx, (chuong, content) in enumerate(chapter_texts.items()):
              amount = distribution[idx]
              print(f"🔄 Sinh {amount} câu hỏi từ chương {chuong}")

              try:
                qs = generate_questions_advance(content, num_questions=amount) ## vừa thay hàm mới ở đây
                print("in ra all_questions ", qs )
                if qs:
                 all_questions.extend(qs)
              except Exception as e:
                 print(f"❌ Lỗi khi sinh câu hỏi từ chương {chuong}: {e}")
                 print("in ra all_questions ", all_questions )

            questions = all_questions
            print("in ra all_questions ", questions )

            
            
            
            
            # đoạn này là kết thúc
       
        #questions = generate_questions(extracted_text)
        #print("Câu hỏi sinh ra:", questions)
        
        if not isinstance(questions, list):
            print("⚠️ Lỗi: generate_questions() không trả về danh sách!", questions)
            return "Lỗi khi sinh câu hỏi! Kiểm tra dữ liệu đầu vào."
        
        # if isinstance(questions, str):
        #     questions = questions.split("\n")
        
        print("📥 Mode nhận được gần đk lưu:", mode)
        
        print("Câu hỏi sinh ra:", all_questions)
        
        print("📥 Mode chính xác nhận được:", repr(mode))

       # print("📌 Dữ liệu JSON đã parse:", json.dumps(data, indent=2, ensure_ascii=False))
        
        if all_questions: # mới fix
            print("📥 Mode chính xác nhận được:", repr(mode))

            if mode == "advanced":
               print("✅ Đang xử lý chế độ ADVANCED")
               # Xử lý dạng nâng cao
               for q in all_questions: # mới fix
                 try:
                   if isinstance(q, dict):
                    question_text = q.get("cau_hoi")
                    answers = q.get("dap_an")
                    correct_answer = q.get("dap_an_dung")
                    chapter = q.get("chuong")
                    do_kho_raw = q.get("do_kho")

                    difficulty = difficulty_map.get(do_kho_raw, None)

                    if not question_text or not answers or not correct_answer:
                        print("⚠️ Câu hỏi thiếu dữ liệu:", q)
                        continue

                    save_ai_generated_question(
                        question=question_text,
                        answers=answers,
                        correct_answer=correct_answer,
                        chapter=chapter,
                        difficulty=difficulty,
                        user_id=user_id,
                        exam_set_id=set_id
                    )
                 except Exception as e:
                  print(f"🚨 Lỗi khi lưu câu hỏi (advanced): {e}")
                  import traceback
                  traceback.print_exc()

            else:
                print("➡️ Đang xử lý chế độ THƯỜNG")
                for q in questions:
                 if isinstance(q, dict) and "question" in q and "options" in q  :
                   correct_answer = q.get("correct_answer") or q.get("answer")
                   if not correct_answer:
                       print("⚠️ Câu hỏi không có đáp án đúng:", q)
                       continue    
                   question_text = q["question"]
                   answers = q["options"]
                   print("🧪 DEBUG - Câu hỏi:", question_text)
                   print("🧪 DEBUG - Đáp án:", answers)
                   print("🧪 DEBUG - Đáp án đúng:", correct_answer)
            
                   try:
                     
                     save_question_to_db(question_text, answers, correct_answer, user_id=user_id,
                        exam_set_id=set_id, chapter=None, difficulty=None )
                   except Exception as e:
                    print(f"🚨 Lỗi khi lưu câu hỏi: {e}")
                   else:
                    print("🚨 Câu hỏi bị lỗi định dạng:", q)  # Chỉ in nếu `q` tồn tại
                 else:
                    print("🚨 Không có câu hỏi nào được tạo!")

            #set_id = get_latest_exam_set_id(user_id=user_id)
            #exam_sets = get_exam_sets_by_user(user_id)
            #print("in ra danh sách bộ đề của người dùng:", exam_sets)
            if not set_id:
                 return "Không có bộ đề nào!", 400
            #return render_template("index.html", questions=questions, exam_sets=exam_sets)
            #return redirect(url_for('question.edit_questions', exam_set_id=set_id))
            return jsonify({"redirect_url": url_for('question.edit_questions', exam_set_id=set_id)})
    upcoming_rooms = get_upcoming_exam_rooms(user_id)

    return render_template("index.html", questions=questions, exam_sets=exam_sets, upcoming_rooms=upcoming_rooms)

    
# def get_db_connection():
#     return mysql.connector.connect(
#         host="localhost",
#         user="root",       # Thay bằng user MySQL của bạn
#         password="",       # Thay bằng password của bạn
#         database="everest" # Đặt tên database
#     )
# ##### HÀM KẾT NỐI CƠ SỞ DỮ LIỆU ĐƯỢC VIẾT LẠI


def get_db_connection():
    return mysql.connector.connect(
        host=os.getenv("MYSQL_HOST", "localhost"),
        user=os.getenv("MYSQL_USER", "root"),
        password=os.getenv("MYSQL_PASSWORD", ""),
        database=os.getenv("MYSQL_DATABASE", "everest")
    )





###### KẾT THÚC VIẾT LẠI
def get_latest_exam_set_id(user_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("""
        SELECT id FROM exam_sets
        WHERE user_id = %s
        ORDER BY created_at DESC LIMIT 1
    """, (user_id,))
    row = cursor.fetchone()
    cursor.close()
    conn.close()
    return row["id"] if row else None
def save_question_to_db(question, answers, correct_answer, user_id, exam_set_id, chapter=None, difficulty=None):
    conn = get_db_connection()
    cursor = conn.cursor()

    query = "INSERT INTO questions (question_text, answer_a, answer_b, answer_c, answer_d, correct_answer,user_id, exam_set_id, chapter, difficulty) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)"
    cursor.execute(query, (question, answers[0], answers[1], answers[2], answers[3], correct_answer, user_id, exam_set_id, chapter, difficulty))

    conn.commit()
    cursor.close()
    conn.close()
def create_exam_set(set_name, user_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    created_at = datetime.now()
    subject_name = "Môn học mặc định"
    source = "from_ai"
    


    cursor.execute("""
        INSERT INTO exam_sets (set_name, user_id, subject_name, created_at, sources)
        VALUES (%s, %s, %s, %s, %s)
    """, (set_name, user_id, subject_name, created_at, source))

    conn.commit()
    exam_set_id = cursor.lastrowid
    cursor.close()
    conn.close()
    return exam_set_id

def save_ai_generated_question(question, answers, correct_answer, user_id, exam_set_id, chapter, difficulty):
    conn = get_db_connection()
    cursor = conn.cursor()

    query = """
    INSERT INTO questions 
    (question_text, answer_a, answer_b, answer_c, answer_d, correct_answer, user_id, exam_set_id, chapter, difficulty)
    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
    """

    cursor.execute(query, (
        question,
        answers.get("A", ""),  # Đáp án A
        answers.get("B", ""),  # Đáp án B
        answers.get("C", ""),  # Đáp án C
        answers.get("D", ""),  # Đáp án D
        correct_answer,         # Đáp án đúng ("A", "B", "C", "D")
        user_id,                # ID người dùng
        exam_set_id,            # ID bộ đề
        chapter,                # Chương
        difficulty              # Mức độ khó (0, 1, 2)
    ))

    conn.commit()
    cursor.close()
    conn.close()



@question_bp.route("/export", methods=["GET"])
def export_to_docx():
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM questions")
    questions = cursor.fetchall()
    cursor.close()
    conn.close()

    # Tạo file DOCX
    doc = Document()
    doc.add_heading("ĐỀ THI TRẮC NGHIỆM", level=1)

    for index, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{index}. {q['question_text']}")
        doc.add_paragraph(f"A. {q['answer_a']}")
        doc.add_paragraph(f"B. {q['answer_b']}")
        doc.add_paragraph(f"C. {q['answer_c']}")
        doc.add_paragraph(f"D. {q['answer_d']}")
        doc.add_paragraph("\n")

    file_path = "exam.docx"
    doc.save(file_path)
    
    return send_file(file_path, as_attachment=True)


# edit questions

@question_bp.route("/questions/edit")
def edit_questions():
    exam_set_id = request.args.get("exam_set_id", type=int)
    user_id = current_user.id  # Mặc định user

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # Kiểm tra bộ đề có thuộc user này không
    cursor.execute("""
        SELECT * FROM exam_sets 
        WHERE id = %s AND user_id = %s
    """, (exam_set_id, user_id))
    exam_sets = cursor.fetchone()

    if not exam_sets:
        cursor.close()
        conn.close()
        return "❌ Bộ đề không hợp lệ hoặc không thuộc quyền của bạn.", 403

    
    cursor.execute("""
        SELECT id, question_text, correct_answer, answer_a, answer_b, answer_c, answer_d
        FROM questions
        WHERE exam_set_id = %s
    """, (exam_set_id,))
    rows = cursor.fetchall()
    
    cursor.execute("""
        SELECT set_name, subject_name
        FROM exam_sets
        WHERE id = %s
    """, (exam_set_id,))
    exam_set = cursor.fetchone()
    
    cursor.close()
    conn.close()

    questions = []
    for q in rows:
        q["answer"] = [q["answer_a"], q["answer_b"], q["answer_c"], q["answer_d"]]
        questions.append(q)

    return render_template("edit_questions.html", user_id=user_id,questions=questions, exam_set_id=exam_set_id, exam_set=exam_set)
# logic edit
@question_bp.route("/questions/update/<int:question_id>", methods=["POST"])
def update_question(question_id):
    user_id = current_user.id
    conn = get_db_connection()
    cursor = conn.cursor()

    # Kiểm tra quyền sửa
    cursor.execute("""
        SELECT q.id FROM questions q
        JOIN exam_sets es ON q.exam_set_id = es.id
        WHERE q.id = %s AND es.user_id = %s
    """, (question_id, user_id))
    check = cursor.fetchone()
    if not check:
        cursor.close()
        conn.close()
        return "❌ Không có quyền sửa câu hỏi này.", 403

    # Lấy dữ liệu từ form
    question_text = request.form.get(f"question_{question_id}")
    correct_answer = request.form.get(f"correct_{question_id}")
    option_a = request.form.get(f"option_{question_id}_0")
    option_b = request.form.get(f"option_{question_id}_1")
    option_c = request.form.get(f"option_{question_id}_2")
    option_d = request.form.get(f"option_{question_id}_3")

    cursor.execute("""
    UPDATE questions 
    SET question_text = %s, 
        answer_a = %s, 
        answer_b = %s, 
        answer_c = %s, 
        answer_d = %s, 
        correct_answer = %s
        WHERE id = %s
    """, (question_text, option_a, option_b, option_c, option_d, correct_answer, question_id))
    conn.commit()

    cursor.close()
    conn.close()
    return redirect(request.referrer or url_for('question.edit_questions'))
# delete
@question_bp.route("/questions/delete/<int:question_id>", methods=["POST"])
def delete_question(question_id):
    user_id = current_user.id
    conn = get_db_connection()
    cursor = conn.cursor()

    # Kiểm tra quyền xóa
    cursor.execute("""
        SELECT q.id FROM questions q
        JOIN exam_sets es ON q.exam_set_id = es.id
        WHERE q.id = %s AND es.user_id = %s
    """, (question_id, user_id))
    check = cursor.fetchone()
    if not check:
        cursor.close()
        conn.close()
        return "❌ Không có quyền xóa câu hỏi này.", 403

    # Xóa
    cursor.execute("DELETE FROM questions WHERE id = %s", (question_id,))
    conn.commit()

    cursor.close()
    conn.close()
    return redirect(request.referrer or url_for('edit_questions'))

## route mới thêm vào
### ROUTE NÀY CẬP NHẬT THÔNG TIN BỘ ĐỀ

@question_bp.route("/update-exam-set-info/<int:exam_set_id>", methods=["POST"])
def update_exam_set_info(exam_set_id):
    user_id = current_user.id
    new_set_name = request.form.get("set_name", "").strip()
    new_subject_name = request.form.get("subject_name", "").strip()

    if not new_set_name or not new_subject_name:
        return jsonify({"success": False, "message": "Tên bộ đề và môn học không được để trống."}), 400

    conn = get_db_connection()
    cursor = conn.cursor()

    # Kiểm tra quyền
    cursor.execute("SELECT user_id FROM exam_sets WHERE id = %s", (exam_set_id,))
    row = cursor.fetchone()

    if not row or row[0] != user_id:
        cursor.close()
        conn.close()
        return jsonify({"success": False, "message": "Bạn không có quyền sửa bộ đề này."}), 403

    # Cập nhật
    cursor.execute("""
        UPDATE exam_sets
        SET set_name = %s, subject_name = %s
        WHERE id = %s
    """, (new_set_name, new_subject_name, exam_set_id))
    
    conn.commit()
    cursor.close()
    conn.close()

    return jsonify({"success": True, "message": "Cập nhật thành công."})


### KẾT THÚC CẬP NHẬT ROUTE






@question_bp.route('/rename-exam-set', methods=['POST'])
@login_required
def rename_exam_set():
    try:
        data = request.get_json()
        exam_set_id = data.get('exam_set_id')
        new_name = data.get('new_name', '').strip()

        if not exam_set_id or not new_name:
            return jsonify({'status': 'error', 'message': 'Thiếu thông tin.'}), 400

        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("""
            UPDATE exam_sets
            SET set_name = %s
            WHERE id = %s AND user_id = %s
        """, (new_name, exam_set_id, current_user.id))

        conn.commit()
        cursor.close()
        conn.close()

        return jsonify({'status': 'success', 'message': 'Cập nhật tên thành công.'})

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500




## kết thúc route này


def get_upcoming_exam_rooms(user_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    now = datetime.now()
    
    
    cursor.execute("""
        SELECT 
            id, room_code, subject_name, duration_minutes, open_time, 
            grace_period_minutes
        FROM exam_rooms
        WHERE created_by_user_id = %s
    """, (user_id,))
    
    all_rooms = cursor.fetchall()
    upcoming_rooms = []

    for room in all_rooms:
        open_time = room['open_time']
        grace = room['grace_period_minutes']
        duration = room['duration_minutes']
        end_time = open_time + timedelta(minutes=grace + duration)

        if end_time > now:
            upcoming_rooms.append(room)

    cursor.close()
    conn.close()

    return upcoming_rooms

    # cursor.execute("""
    #     SELECT room_code, subject_name, open_time, duration_minutes
    #     FROM exam_rooms
    #     WHERE created_by_user_id = %s AND open_time > %s
    #     ORDER BY open_time ASC
    # """, (user_id, now))

    # rooms = cursor.fetchall()

    # cursor.close()
    # conn.close()

    # return rooms

def get_exam_sets_by_user(user_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("""
        SELECT id, set_name, created_at 
        FROM exam_sets 
        WHERE user_id = %s 
        ORDER BY created_at DESC
    """, (user_id,))
    exam_sets = cursor.fetchall()
    cursor.close()
    conn.close()
    return exam_sets

def is_full_thpt2025_exam_set(exam_set_id):
    conn = get_db_connection()
    cursor = conn.cursor()

    # Cần kiểm tra cả hai bảng đều có dữ liệu
    tf_query = "SELECT 1 FROM tf_questions WHERE exam_set_id = %s LIMIT 1"
    sa_query = "SELECT 1 FROM short_answer_questions WHERE exam_set_id = %s LIMIT 1"

    cursor.execute(tf_query, (exam_set_id,))
    has_tf = cursor.fetchone() is not None

    cursor.execute(sa_query, (exam_set_id,))
    has_sa = cursor.fetchone() is not None

    cursor.close()
    conn.close()

    return has_tf and has_sa
def get_mcq_original_questions_by_exam_set(exam_set_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    if is_advanced_mcq_exam_set(exam_set_id):
        # MCQ nâng cao: cần nhiều trường hơn
        query = """
            SELECT id, question_text, difficulty, chapter,
                   answer_a, answer_b, answer_c, answer_d, correct_answer
            FROM questions
            WHERE exam_set_id = %s
            ORDER BY id
        """
    else:
        # MCQ thuần: chỉ cần một số trường
        query = """
            SELECT id, question_text, answer_a, answer_b, answer_c, answer_d, correct_answer
            FROM questions
            WHERE exam_set_id = %s
            ORDER BY id
        """

    cursor.execute(query, (exam_set_id,))
    questions = cursor.fetchall()

    cursor.close()
    conn.close()
    #print("in ra questions", questions)
    return questions
def is_advanced_mcq_exam_set(exam_set_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    query = """
        SELECT COUNT(*) FROM questions
        WHERE exam_set_id = %s AND (difficulty IS NOT NULL OR chapter IS NOT NULL)
    """
    
    cursor.execute(query, (exam_set_id,))
    result = cursor.fetchone()
    
    cursor.close()
    conn.close()
    
    return result[0] > 0
def get_original_questions_thpt2025(exam_set_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # Lấy MCQ
    cursor.execute("""
        SELECT id, 'MCQ' AS type, question_text, answer_a, answer_b, answer_c, answer_d, correct_answer
        FROM questions
        WHERE exam_set_id = %s
        ORDER BY id
    """, (exam_set_id,))
    mcq = cursor.fetchall()

    # Lấy TF
    cursor.execute("""
        SELECT tf.id, 'tf' AS question_type, tf.question_text, ti.statement_text, ti.is_true, tf.image_url
        FROM tf_questions tf
        JOIN tf_items ti ON tf.id = ti.tf_question_id
        WHERE tf.exam_set_id = %s
        ORDER BY tf.id, ti.id
    """, (exam_set_id,))
    tf_raw = cursor.fetchall()

    tf_dict = {}
    for item in tf_raw:
        qid = item['id']
        if qid not in tf_dict:
            tf_dict[qid] = {
                'question_type': 'tf',
                'question_text': item['question_text'],
                'image_url': item['image_url'],  # ✅ Thêm dòng này
                'sub_items': []
            }
        tf_dict[qid]['sub_items'].append({
            'statement_text': item['statement_text'],
            'is_true': item['is_true']
        })
    tf = list(tf_dict.values())

    # Lấy SA
    cursor.execute("""
        SELECT id, 'SA' AS question_type, question_text, correct_answer
        FROM short_answer_questions
        WHERE exam_set_id = %s
        ORDER BY id
    """, (exam_set_id,))
    sa = cursor.fetchall()

    cursor.close()
    conn.close()

    return {
        'MCQ': mcq,
        'tf': tf,
        'SA': sa
    }
def has_exam_codes(exam_set_id):
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) FROM exam_codes WHERE exam_set_id = %s", (exam_set_id,))
    result = cursor.fetchone()[0]
    cursor.close()
    conn.close()
    return result > 0

def get_exam_code_data(exam_code_id):
    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    # Lấy thông tin mã đề và exam_set_id
    cursor.execute("""
        SELECT ec.id AS exam_code_id, ec.code AS exam_code, ec.exam_set_id
        FROM exam_codes ec
        WHERE ec.id = %s
    """, (exam_code_id,))
    code_info = cursor.fetchone()
    if not code_info:
        cursor.close()
        conn.close()
        return None

    exam_set_id = code_info['exam_set_id']

    # -------- Lấy câu hỏi MCQ từ exam_question_map ----------
    cursor.execute("""
        SELECT q.id AS question_id,
               q.question_text,
               eqm.answer_a,
               eqm.answer_b,
               eqm.answer_c,
               eqm.answer_d,
               eqm.correct_answer,
               eqm.question_order,
               eqm.chapter,
               eqm.difficulty
        FROM exam_question_map eqm
        JOIN questions q ON q.id = eqm.question_id
        WHERE eqm.exam_code_id = %s
        ORDER BY eqm.question_order
    """, (exam_code_id,))
    mcq_questions = cursor.fetchall()

    # --------- Kiểm tra xem có TF hoặc SA không ------------
    cursor.execute("""
        SELECT hqm.id, hqm.question_type, hqm.original_question_id, hqm.position
        FROM hs_question_map hqm
        WHERE hqm.exam_code_id = %s
        ORDER BY hqm.position
    """, (exam_code_id,))
    hs_mappings = cursor.fetchall()

    tf_questions = []
    sa_questions = []

    for m in hs_mappings:
        if m['question_type'] == 'tf':
            cursor.execute("SELECT * FROM tf_questions WHERE id = %s", (m['original_question_id'],))
            tf_q = cursor.fetchone()
            if tf_q:
                cursor.execute("""
                    SELECT label, content, is_true, position
                    FROM hs_tf_statements
                    WHERE hs_map_id = %s
                    ORDER BY position
                """, (m['id'],))
                statements = cursor.fetchall()
                tf_q['statements'] = statements
                tf_questions.append({'position': m['position'], **tf_q})

        elif m['question_type'] == 'sa':
            cursor.execute("SELECT * FROM short_answer_questions WHERE id = %s", (m['original_question_id'],))
            sa_q = cursor.fetchone()
            if sa_q:
                sa_questions.append({'position': m['position'], **sa_q})

    cursor.close()
    conn.close()

    # Xác định loại đề
    if tf_questions or sa_questions:
        question_type = "thpt_2025"
    else:
        question_type = "mcq"

    return {
        "exam_code": code_info['exam_code'],
        "exam_code_id": code_info['exam_code_id'],
        "exam_set_id": code_info['exam_set_id'],
        "question_type": question_type,
        "mcq_questions": mcq_questions,
        "tf_questions": tf_questions,
        "sa_questions": sa_questions
    }
    
@question_bp.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    user_id = current_user.id
    if not user_id:
        flash("Bạn chưa đăng nhập.")
        return redirect(url_for("auth.login"))

    conn = get_db_connection()
    cursor = conn.cursor(dictionary=True)

    cursor.execute("SELECT username, email, role, password FROM users WHERE id = %s", (user_id,))
    user = cursor.fetchone()

    if request.method == 'POST':
        email = request.form.get('email')
        old_password = request.form.get('old_password')
        new_password = request.form.get('new_password')

        # ✅ Cập nhật email
        if email:
            cursor.execute("UPDATE users SET email = %s WHERE id = %s", (email, user_id))

        # ✅ Nếu người dùng muốn đổi mật khẩu
        if new_password:
            if not old_password:
                flash("Vui lòng nhập mật khẩu cũ để đổi mật khẩu.")
                return redirect(url_for("user.profile"))

            if not check_password_hash(user["password"], old_password):
                flash("Mật khẩu cũ không đúng.")
                return redirect(url_for("user.profile"))

            hashed = generate_password_hash(new_password)
            cursor.execute("UPDATE users SET password = %s WHERE id = %s", (hashed, user_id))

        conn.commit()
        flash("Cập nhật thông tin thành công!")
        return redirect(url_for("user.profile"))

    cursor.close()
    conn.close()

    return render_template("profile.html", user=user)









# kết thúc hàm lấy danh sách bộ đề






# if __name__ == "__main__":
#     app.run(debug=True)
